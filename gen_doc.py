# -*- coding: utf-8 -*-
"""
Lumina 产品介绍书 —— 营销部门出品
面向高校的知识型社区 App · 后端技术深度解读
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

os.chdir('C:/Lumina')
doc = Document()

# ---- 全局样式 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.6
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.8)
    s.right_margin = Cm(2.8)

# ---- 工具函数 ----
def h(text, level=1):
    """标题"""
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def p(text, bold=False, color=None, size=11, align=None):
    """正文段落"""
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.font.size = Pt(size)
    r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    pp.paragraph_format.space_after = Pt(6)
    pp.paragraph_format.line_spacing = 1.6
    if align: pp.alignment = align
    return pp

def b(text, level=0):
    """项目符号"""
    pp = doc.add_paragraph(text, style='List Bullet')
    pp.paragraph_format.space_after = Pt(3)
    pp.paragraph_format.line_spacing = 1.5
    for run in pp.runs:
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level > 0:
        pp.paragraph_format.left_indent = Cm(1.2 * level)
    return pp

def feat(name, desc):
    """特性条目：蓝色标题 + 描述"""
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(4)
    pp.paragraph_format.line_spacing = 1.5
    r1 = pp.add_run('▸ ' + name + '：')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0, 102, 255)
    r1.font.name = '微软雅黑'
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = pp.add_run(desc)
    r2.font.size = Pt(10.5); r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def code(text):
    """代码/端点标注"""
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(2)
    pp.paragraph_format.line_spacing = 1.3
    r = pp.add_run('  ' + text)
    r.font.name = 'Consolas'; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(70, 70, 70)

def divider():
    """分隔线"""
    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after = Pt(4)
    r = pp.add_run('─' * 50)
    r.font.color.rgb = RGBColor(200, 200, 200)
    r.font.size = Pt(8)

def info_box(title, lines):
    """信息框：用缩进段落模拟"""
    p('▌ ' + title, bold=True, color=(0, 102, 255), size=12)
    for line in lines:
        b(line)

# ====================================================================
#                          封  面
# ====================================================================
for _ in range(4):
    doc.add_paragraph()

t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run('LUMINA')
r.font.size = Pt(58); r.bold = True; r.font.color.rgb = RGBColor(0, 102, 255)
r.font.name = 'Georgia'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

t1b = doc.add_paragraph(); t1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1b = t1b.add_run('知识型学术社区 · 产品与技术白皮书')
r1b.font.size = Pt(16); r1b.font.color.rgb = RGBColor(100, 100, 100)
r1b.font.name = '微软雅黑'
r1b.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('以知识为星尘，让深度思考被看见')
r2.font.size = Pt(14); r2.italic = True
r2.font.color.rgb = RGBColor(0, 102, 255)
r2.font.name = '微软雅黑'
r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

today = datetime.date.today().strftime('%Y 年 %m 月 %d 日')
meta_lines = [
    f'版本：v1.0.0       日期：{today}',
    '技术栈：HarmonyOS 5.0 (ArkTS) + Node.js (Express)',
    '适用范围：华为手机 / 平板 · API 12+',
]
for ml in meta_lines:
    tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = tt.add_run(ml); rr.font.size = Pt(10.5)
    rr.font.color.rgb = RGBColor(130, 130, 130)
    rr.font.name = '微软雅黑'
    rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ====================================================================
#                          目  录
# ====================================================================
h('目  录', 1)
toc_items = [
    ('第一章', '产品定位与市场分析'),
    ('第二章', '核心功能全景'),
    ('第三章', '后端系统技术架构 ★'),
    ('第四章', 'API 接口详细设计'),
    ('第五章', '安全体系与数据保护'),
    ('第六章', '前端架构与体验设计'),
    ('第七章', '创新特性深度解析'),
    ('第八章', '部署运维与演示方案'),
    ('第九章', '技术成果与迭代展望'),
]
for ch, title in toc_items:
    pp = doc.add_paragraph()
    r_ch = pp.add_run(f'{ch}  ')
    r_ch.bold = True; r_ch.font.size = Pt(12); r_ch.font.color.rgb = RGBColor(0, 102, 255)
    r_ch.font.name = '微软雅黑'
    r_ch.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r_ti = pp.add_run(title)
    r_ti.font.size = Pt(12); r_ti.font.name = '微软雅黑'
    r_ti.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pp.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ====================================================================
#              第一章：产品定位与市场分析
# ====================================================================
h('第一章  产品定位与市场分析', 1)

h('1.1  产品愿景', 2)
p('在这个信息过载的时代，高校学生群体长期面临一个核心矛盾：一方面，他们渴望获取有深度、有质量的知识内容来进行自我提升；另一方面，主流社交平台的内容生态日益碎片化和娱乐化，真正有价值的学术讨论和深度见解反而被淹没在短视频和信息流之中。')
p('Lumina 正是为了解决这一矛盾而生。我们定位为一款面向高校学生的"知识型学术社区"——它不是一个冷冰冰的问答工具，也不是一个泛娱乐化的社交平台，而是一个以知识分享为纽带、以优质内容为核心驱动力的学术社交空间。')
p('我们的品牌理念凝结为一句 Slogan："以知识为星尘，照亮彼此的视野"。在这句话背后，是一整套从产品设计到技术实现的完整思考。', bold=True, color=(0, 102, 255))

h('1.2  目标用户画像', 2)
b('核心用户：在校本科生及研究生（18-26 岁），对学术讨论、科技前沿、人文社科、职业发展等话题有深度兴趣')
b('内容生产者：乐于分享学习笔记、课程心得、技术教程、读书感悟的"知识输出型"学生')
b('内容消费者：希望通过高质量信息流获取知识和见解，而非消磨时间的"成长型"学生')
b('社区运营者：未来可扩展至高校教师、实验室课题组、学术社团等机构用户')

h('1.3  市场差异化', 2)
p('与现有竞品的核心差异：')

# 对比表格
table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['维度', '传统社交平台', '传统问答平台', 'Lumina']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for pp in cell.paragraphs:
        for run in pp.runs:
            run.bold = True; run.font.size = Pt(10)

data = [
    ['内容调性', '娱乐化、碎片化', '工具化、冷冰冰', '知识型、有温度的学术社区'],
    ['激励方式', '点赞数/粉丝数', '积分/徽章', '星尘虚拟货币+六等级声望体系'],
    ['用户关系', '单向关注', '匿名问答', '双向关注+知识打赏经济闭环'],
    ['内容质量', '算法推荐驱动', '搜索驱动', '社区打赏驱动+热榜算法驱动'],
    ['平台', 'Android/iOS 通用', 'Web 为主', 'HarmonyOS 原生，深度适配华为生态'],
]
for r_idx, row_data in enumerate(data):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(9.5)

doc.add_paragraph()

h('1.4  HarmonyOS 生态战略', 2)
p('Lumina 选择 HarmonyOS 作为首发平台，不仅是技术选型，更是战略选择。华为 HarmonyOS 在国内高校学生群体中拥有快速增长的用户基础，而鸿蒙生态中的优质原生应用——尤其是面向学术场景的应用——仍然相对稀缺。Lumina 作为一款鸿蒙原生 App，能够在应用市场中获得差异化的曝光优势，同时深度利用鸿蒙的分布式能力（如同账号多设备同步）为未来的跨设备体验奠定基础。')

doc.add_page_break()

# ====================================================================
#              第二章：核心功能全景
# ====================================================================
h('第二章  核心功能全景', 1)

h('2.1  功能架构总览', 2)
p('Lumina 的功能体系围绕"内容生产 → 内容消费 → 社区互动 → 激励正反馈"这一核心链路展开，形成四个功能象限：')

info_box('第一象限：账号与社交关系',
         ['学号注册 / 密码登录（PBKDF2-SHA512 加密）',
          '个人资料编辑（昵称、简介、头像、院校信息）',
          '关注 / 取消关注（双向计数实时同步）',
          '粉丝列表 & 关注列表（支持跳转个人主页）'])

info_box('第二象限：内容社区',
         ['帖子发布（标题 + 正文 + 10 种知识标签 + 最多 9 张图片）',
          '首页信息流（时间排序 + 标签分类筛选）',
          '热榜系统（加权热度算法：点赞×2 + 评论×3 + 星尘×5 + 浏览×0.1）',
          '帖子互动（点赞/评论/打赏，三者独立计数）',
          '全字段搜索（标题 + 正文 + 标签模糊匹配）'])

info_box('第三象限：星尘激励体系',
         ['每日签到领星尘（+5/天）',
          '星尘打赏机制（6 档金额 + 自定义留言）',
          '六等级晋升系统（星尘新芽 → 星海领主）',
          '星尘收支明细（完整流水记录）'])

info_box('第四象限：消息与通知',
         ['分类通知（全部/赞/评论/官方/私信）',
          '官方消息广播（管理员推全站）',
          '未读角标（底部 Tab 红点实时更新）',
          '私信聊天（会话列表 + 一对一聊天）'])

h('2.2  页面导航地图', 2)
p('应用共包含 20 个独立页面，形成清晰的导航层级：')
b('启动层：SplashPage（品牌闪屏 → 2 秒后自动跳转主页）')
b('主页层：Index（底部 5 个 Tab：首页/热榜/[发布]/消息/我的）')
b('内容层：HomePage / HotPage / PostDetailPage / PublishPage / SearchPage')
b('社交层：ChatListPage / ChatPage / FollowListPage / UserProfilePage')
b('用户层：LoginPage / RegisterPage / ProfilePage / EditProfilePage / SettingsPage')
b('系统层：AboutPage / StarDustRewardPage / StarDustRulesPage')

doc.add_page_break()

# ====================================================================
#         第三章：后端系统技术架构 ★（核心重点章节）
# ====================================================================
h('第三章  后端系统技术架构 ★', 1)

p('本章是 Lumina 产品技术白皮书的核心章节。我们将从技术选型决策、服务器架构设计、数据模型建模、API 设计哲学、安全防护体系、通知引擎实现等六个维度，深度解析 Lumina 后端系统的完整技术架构。', bold=True)

# ----- 3.1 -----
h('3.1  技术选型决策分析', 2)
p('后端技术栈的选型并非随意决定，而是经过对项目需求的深入分析后做出的工程决策。以下是我们的选型逻辑：')

feat('Node.js + Express 框架',
     '选择 Express 而非更重的 Nest.js 或 Koa，是基于"课设阶段快速迭代 + 代码可读性"的双重考量。Express 是 Node.js 生态中最成熟、社区最大的 Web 框架——GitHub 60k+ stars，npm 周下载量 3000 万+。它的中间件模型（middleware pattern）直观易懂，路由定义清晰，非常适合 RESTful API 的快速构建。同时，选择 JavaScript 而非 TypeScript 编写后端，是为了让同一门语言（ArkTS 是 TS 超集）贯穿前后端，降低全栈开发的心智负担。')
feat('JSON 文件数据库',
     '对于课设阶段的 MVP（最小可行产品），使用 db.json 而非 MySQL/PostgreSQL/MongoDB 是一个务实的工程决策。原因有三：其一，JSON 文件无需安装和配置数据库服务，做到"clone 即运行"；其二，文件数据库天然支持版本控制（纳入 Git），答辩时可直接用文本编辑器打开展示数据结构；其三，所有数据操作本质上是对原生 JavaScript 对象的读写，不存在 ORM 映射损耗。当未来需要迁移到生产级数据库时，由于 API 层已经完整抽象，替换底层存储仅需修改 readDB/writeDB 两个函数。')
feat('仅两个核心依赖',
     'package.json 中只声明了 express 和 cors 两个核心依赖。身份认证的 Token 生成和密码哈希均使用 Node.js 内置的 crypto 模块实现，文件读写使用内置的 fs 模块，零第三方认证库依赖。这种极简依赖策略不仅减少了供应链安全风险（无 left-pad 类问题），也展示了从底层理解安全机制的工程能力。')

# ----- 3.2 -----
h('3.2  服务器架构设计', 2)
p('Lumina 后端采用"单文件模块化"架构——所有业务逻辑集中在 server.js 一个文件中，但通过清晰的代码分区（认证模块、帖子模块、评论模块、星尘模块、用户模块、通知模块、关注模块、私信模块、搜索模块、管理模块）实现了逻辑上的模块化。')

p('架构层次（从外到内）：', bold=True)
b('第 1 层 - 中间件链：CORS 跨域 → JSON 解析 → 认证中间件（Token 验证）')
b('第 2 层 - 路由分发：Express Router 将 25 个端点按 URL 路径分发到对应的处理函数')
b('第 3 层 - 业务逻辑：每个端点处理函数中实现参数校验、数据读写、业务规则判断')
b('第 4 层 - 数据持久化：readDB() / writeDB() 两个函数封装对 db.json 的全部 IO 操作')
b('第 5 层 - 基础设施：genId() 唯一 ID 生成、hashPassword() 密码哈希、setToken() Token 管理')

p('')
p('关键设计原则：', bold=True)
b('RESTful 语义：遵循 HTTP 方法语义——GET 读取、POST 创建、PUT 更新。URL 路径以资源名词命名（/api/posts、/api/users），避免动词式命名')
b('统一响应格式：所有接口返回 { code: 0, data: ..., message: ... } 格式。code=0 表示成功，code=1 表示业务错误。前端只需检查 code 字段即可判断请求结果')
b('无状态认证：采用 Bearer Token 方案，服务端内存 Map 存储 Token 映射，不依赖 Session/Cookie')
b('原子化数据操作：每个请求处理函数中 readDB() → 修改 → writeDB() 形成一个原子操作单元')

# ----- 3.3 -----
h('3.3  数据模型设计', 2)
p('Lumina 的 JSON 数据库设计了 8 个数据集合，对应关系型数据库中的 8 张表。以下逐一分析每个集合的设计意图和关键字段：')

h('3.3.1  用户集合 (users)', 3)
p('用户是整个系统的核心实体。除了基本的身份信息外，我们在用户模型中内嵌了星尘余额、等级、社交计数等字段，避免了高频查询时的多表 JOIN 操作。')
code('关键字段：id, studentId, nickname, avatar, bio, campus, major, grade')
code('安全字段：passwordHash (PBKDF2-SHA512), passwordSalt (16字节随机盐)')
code('星尘字段：starDustBalance, totalStarDustReceived, level (1-6), levelTitle')
code('统计字段：postCount, followingCount, followerCount')
code('设计考量：将社交计数（postCount/followingCount/followerCount）冗余存储在用户表中，')
code('          避免每次展示用户卡片时都需要 COUNT 关联表，这是一个典型的"空间换时间"优化。')

h('3.3.2  帖子集合 (posts)', 3)
code('关键字段：id, title, content, images[], authorId, authorName, tag, createTime')
code('互动字段：likeCount, commentCount, starDustCount, viewCount, likedBy[]')
code('状态字段：isTop (置顶), isHot (热门), campusId (校区)')
code('设计考量：likedBy 数组存储所有点赞用户 ID，以 O(n) 查找替代独立点赞表，简化数据模型。')
code('          authorName 冗余存储避免展示帖子时必须 JOIN 用户表。')

h('3.3.3  评论集合 (comments)', 3)
code('关键字段：id, postId, content, authorId, authorName, createTime, likeCount, likedBy[]')
code('关联方式：通过 postId 外键关联帖子，查询时在前端按 createTime 降序排列。')

h('3.3.4  星尘记录集合 (starDustRecords)', 3)
code('关键字段：id, fromUserId, fromUserName, toUserId, toUserName, postId, amount, message')
code('设计意图：完整记录每一笔星尘流转，支持用户查询收支明细。这是星尘经济系统的"账本"。')

h('3.3.5  通知集合 (notifications)', 3)
code('关键字段：id, type (like|comment|reward|official), fromUserId, toUserId, postId')
code('状态字段：isRead (标记已读/未读), createTime')
code('设计考量：type 字段区分 4 种通知类型，前端按 type 实现分类 Tab 筛选。')
code('          official 类型的通知通过 /api/admin/broadcast 接口批量创建。')

h('3.3.6  关注集合 (follows)', 3)
code('关键字段：fromUserId, toUserId, createTime')
code('设计考量：以 fromUserId + toUserId 二元组唯一标识一条关注关系。')
code('          查询粉丝数 = follows.filter(toUserId == id).length。')

h('3.3.7  会话集合 (conversations)', 3)
code('关键字段：id, user1, user2, lastMessage, lastTime')
code('设计考量：user1 和 user2 无顺序保证，查找时需双向匹配。lastMessage 冗余存储最后一条消息预览。')

h('3.3.8  消息集合 (messages)', 3)
code('关键字段：id, conversationId, fromUserId, toUserId, content, createTime, isRead')

doc.add_page_break()

# ----- 3.4 -----
h('3.4  中间件与请求处理管道', 2)
p('Lumina 的请求处理管道由三层中间件组成，每一层负责特定职责：')

feat('第 1 层 - CORS 跨域中间件',
     'app.use(cors())。允许来自任意来源的 HTTP 请求。在开发阶段，HarmonyOS 模拟器通过 10.0.2.2 访问宿主机，属于跨域请求。生产环境应限制为特定域名。')
feat('第 2 层 - JSON 解析中间件',
     'app.use(express.json())。自动将请求体中的 JSON 字符串解析为 JavaScript 对象，挂载到 req.body。支持 UTF-8 编码的中文内容。')
feat('第 3 层 - 认证中间件 (authMiddleware)',
     '这是 Lumina 后端的核心安全组件。中间件从 HTTP Header 中提取 Authorization: Bearer <token>，在服务端的 tokenStore Map 中查找对应的 userId，并将 userId 注入到 req 对象中供下游处理函数使用。关键设计：认证中间件不拦截未登录请求——它只是"尝试解析"Token，解析不到时 req.userId 保持 undefined，由各业务端点自行判断是否需要登录。这种"软认证"设计保持了中间件的通用性。')

h('3.5  热榜算法设计', 2)
p('Lumina 的热榜排序不是简单的按点赞数降序，而是实现了一个加权热度公式：')
p('热度分数 = 点赞数 × 2 + 评论数 × 3 + 星尘数 × 5 + 浏览数 × 0.1', bold=True, size=12, color=(0, 102, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
p('')
p('加权逻辑反映了我们对不同互动行为"价值权重"的判断：星尘打赏权重最高（5x），因为打赏者付出了自己的虚拟资产，代表了最强的内容认可信号；评论权重次之（3x），因为评论需要用户主动输入文字，参与度高于简单点赞；点赞权重为 2x，是门槛最低的互动；浏览权重仅 0.1x，避免"标题党"通过高点击量刷榜。')
p('这一算法在服务端执行——每次 GET /api/posts/hot 请求时，对全站帖子实时计算热度并降序排列。对于课设阶段的数据量（<1000 条帖子），实时计算完全在性能可接受范围内。未来数据量增长时，可引入定时任务预计算 + Redis 缓存热榜结果。')

# ----- 3.6 -----
h('3.6  ID 生成策略', 2)
p('Lumina 使用自定义的 genId(prefix) 函数生成全局唯一 ID。ID 格式为：{前缀}_{时间戳36进制}{随机字符串}。例如 p_mqs88244mfrh（帖子 ID）、u_mqs7y8rk43vo（用户 ID）。')
p('选择自研 ID 生成而非 UUID 或自增 ID 的原因：前缀让 ID 具有"自描述性"（从 ID 即可判断实体类型）；时间戳部分确保大体有序；随机部分提供足够的唯一性保证。在单机课设场景下，这一策略简单可靠。')

doc.add_page_break()

# ====================================================================
#              第四章：API 接口详细设计
# ====================================================================
h('第四章  API 接口详细设计', 1)

p('Lumina 后端共实现 25 个 RESTful API 端点，按业务模块分为 7 组。所有接口遵循统一规范：URL 以 /api/ 为前缀；请求体使用 JSON 格式 (Content-Type: application/json)；响应体统一为 { code, data, message } 三元组。', bold=True)

h('4.1  认证模块（3 个端点）', 2)
feat('POST /api/auth/register',
     '用户注册。入参：{ studentId, password, nickname }。处理流程：校验学号唯一性 → 生成 16 字节随机盐 → PBKDF2-SHA512 1000 次迭代哈希 → 创建用户对象（含初始 100 星尘） → 生成 Token → 返回 Token + 用户信息。业务规则：密码最短 6 位，学号不可重复。')
feat('POST /api/auth/login',
     '用户登录。入参：{ studentId, password }。处理流程：学号查找用户 → 取出 salt → 对输入密码重新哈希 → 比对哈希值 → 生成 Token → 返回 Token + 脱敏用户信息。密码比对使用 crypto.timingSafeEqual 的等效逻辑，防止时序攻击。')
feat('GET /api/auth/me',
     '获取当前用户。依赖 authMiddleware 解析出的 req.userId。使用解构赋值排除 passwordHash 和 passwordSalt 后返回，确保敏感数据不泄露。')

h('4.2  用户模块（10 个端点）', 2)
code('GET    /api/users/:id                获取用户信息（含等级、星尘余额、社交统计）')
code('PUT    /api/users/:id                更新用户资料（支持部分更新，未传字段保持不变）')
code('POST   /api/users/:id/checkin        每日签到（+5 星尘到余额）')
code('POST   /api/users/:id/follow         关注用户（校验不能自关注、不能重复关注）')
code('POST   /api/users/:id/unfollow       取消关注（同步扣减双方关注/粉丝计数）')
code('GET    /api/users/:id/follow-status  查询关注状态（?userId=当前用户ID）')
code('GET    /api/users/:id/followers      粉丝列表（返回昵称、头像、简介）')
code('GET    /api/users/:id/following      关注列表（同上）')
code('GET    /api/users/:id/posts          用户帖子列表（按时间降序）')
code('GET    /api/users/:id/rewards        星尘收支记录（双向匹配：发出或收到）')

h('4.3  帖子模块（8 个端点）', 2)
code('GET    /api/posts                     帖子列表（?tag=分类筛选，置顶优先→时间降序）')
code('POST   /api/posts                     发布帖子（发帖者 +2 星尘奖励，postCount+1）')
code('GET    /api/posts/hot                 热榜（加权热度实时计算降序）')
code('GET    /api/posts/:id                 帖子详情（浏览量自动 +1）')
code('POST   /api/posts/:id/like           点赞/取消点赞 Toggle（检查 likedBy 数组，存在则移除、不存在则添加，同时创建通知）')
code('GET    /api/posts/:id/comments        评论列表（按时间降序）')
code('POST   /api/posts/:id/comments        发表评论（commentCount+1，创建通知）')
code('POST   /api/posts/:id/reward          星尘打赏（余额校验 → 扣除 → 增加 → 记录 → 等级升级判定 → 通知）')

h('4.4  通知模块（3 个端点）', 2)
code('GET    /api/notifications             获取通知（?userId=当前用户，4 类混排按时间降序）')
code('GET    /api/notifications/unread-count 未读计数（用于底部 Tab 角标）')
code('POST   /api/notifications/read        标记已读（传 notificationId 单条已读，只传 userId 全部已读）')

h('4.5  消息模块（3 个端点）', 2)
code('POST   /api/messages/send             发送私信（自动查找或创建会话，更新 lastMessage）')
code('GET    /api/messages/conversations    会话列表（?userId=，含未读计数和对方信息）')
code('GET    /api/messages/:convId          会话消息（按时间正序，自动标记已读）')

h('4.6  搜索模块（1 个端点）', 2)
code('GET    /api/search?q=关键字           全字段模糊匹配（标题+正文+标签 toLowerCase 包含判断）')

h('4.7  管理模块（4 个端点）', 2)
code('POST   /api/admin/seed                生成 8 篇知识型种子帖子（含真实学术内容）')
code('POST   /api/admin/bootstrap           一键初始化演示数据（6 粉丝 + 关注 + 点赞评论 + 官方消息）')
code('POST   /api/admin/broadcast           全站广播（向所有用户推送官方消息通知）')
code('POST   /api/admin/generate-followers  批量生成粉丝（指定目标用户和数量）')

doc.add_page_break()

# ====================================================================
#              第五章：安全体系与数据保护
# ====================================================================
h('第五章  安全体系与数据保护', 1)

p('在一个面向高校学生的社交平台中，用户数据安全——尤其是账号密码的保护——是不可妥协的底线。Lumina 从零开始实现了完整的后端安全体系，以下逐一分析。', bold=True)

h('5.1  密码安全：PBKDF2-SHA512 加盐哈希', 2)
p('Lumina 绝不以明文存储用户密码。我们采用了业界标准的 PBKDF2 (Password-Based Key Derivation Function 2) 算法，配合 SHA-512 哈希函数和 1000 次迭代，对每个用户的密码进行独立的加盐哈希处理。')
p('完整流程：')
b('用户注册时：调用 crypto.randomBytes(16) 生成 16 字节（128 位）密码学安全随机盐值')
b('使用 crypto.pbkdf2Sync(password, salt, 1000, 64, \'sha512\') 生成 64 字节哈希值')
b('将哈希值和盐值分别存储在用户记录的 passwordHash 和 passwordSalt 字段中')
b('用户登录时：取出该用户的 salt → 对输入的密码重新执行相同哈希 → 比对哈希值')
p('这一设计的核心安全优势在于：即使两个用户设置了完全相同的密码，由于各自拥有独立的随机盐值，数据库中存储的哈希值也完全不同。即使攻击者获取了 db.json 文件，也无法通过彩虹表（Rainbow Table）反推原始密码——他必须对每个用户的盐值单独进行暴力枚举，这在实际操作中是不可行的。')

h('5.2  Token 认证机制', 2)
p('Lumina 实现了轻量级的 Token-Based 认证，不依赖 JWT 库或 OAuth 框架：')
b('Token 格式：lm_{时间戳36进制}_{随机字符串16位}。lm 前缀标识 Lumina 系统，中间段为时间戳的 36 进制表示，末尾为 16 位随机字符')
b('存储方式：服务端内存 Map<token, userId>。每个用户登录时，先清除该用户的所有旧 Token（保证一个用户只有一个有效 Token），再生成新 Token')
b('传递方式：客户端在 HTTP Header 中携带 Authorization: Bearer <token>')
b('安全特性：Token 不含任何用户信息（与 JWT 不同），即使泄露也无法解码出用户数据。Token 与服务端状态绑定，可随时通过重启服务实现全局登出')

h('5.3  数据脱敏', 2)
p('在所有返回用户数据的 API 端点中，后端使用 ES6 解构赋值语法排除敏感字段：')
code('const { passwordHash, passwordSalt, ...safeUser } = user;')
code('res.json({ code: 0, data: safeUser });')
p('这一模式确保 passwordHash 和 passwordSalt 永远不会在 API 响应中出现。即使未来新增敏感字段，也只需在解构语句中增加一个排除项即可。')

h('5.4  业务安全校验', 2)
b('学号唯一性：register 端点检查学号是否已被注册，防止重复注册')
b('密码强度：注册时检查密码长度 >= 6 位')
b('自关注防护：follow 端点校验 userId !== targetId')
b('重复关注防护：follow 端点检查关注关系是否已存在')
b('余额校验：打赏端点检查 fromUser.starDustBalance >= amount，余额不足时拒绝交易')
b('通知自屏蔽：createNotification 函数跳过 fromUserId === toUserId 的情况（不给自己发通知）')

doc.add_page_break()

# ====================================================================
#              第六章：前端架构与体验设计
# ====================================================================
h('第六章  前端架构与体验设计', 1)

h('6.1  技术栈', 2)
b('平台：HarmonyOS 5.0 (API 12)，兼容 Phone + Tablet')
b('语言：ArkTS（华为自研 TypeScript 超集，完整兼容 TS 语法）')
b('UI 范式：ArkUI 声明式开发（组件树 + 状态驱动渲染，类似 SwiftUI / Flutter）')
b('构建系统：Hvigor 5.0（HarmonyOS 原生增量构建工具）')
b('IDE：DevEco Studio 6.0（基于 IntelliJ IDEA Community Edition）')

h('6.2  项目架构分层', 2)
p('Lumina 前端采用六层架构：')
b('Entry 层：EntryAbility.ets — App 生命周期（onCreate/onDestroy/onForeground/onBackground）+ 窗口管理')
b('Page 层：20 个 @Entry 装饰的页面组件，各自对应独立路由')
b('Component 层：PostCard 可复用帖子卡片，被 HomePage/HotPage/SearchPage 等多处引用')
b('Service 层：ApiClient（HTTP 封装 + 自动 Token 注入）+ PostApi + UserApi（业务 API 封装）')
b('Model 层：PostModel / UserModel / CommentModel / StarDustModel（JSON → 类实例安全映射 + 数据格式化方法）')
b('Theme 层：ThemeManager + C 颜色代理类（约 150 行，支持深色/浅色双主题动态切换）')

h('6.3  状态管理机制', 2)
p('Lumina 利用 HarmonyOS AppStorage 实现了跨组件的响应式状态同步：')
feat('loginVersion 全局版本号',
     '登录或退出时更新 loginVersion = Date.now()。所有关心登录状态的组件通过 @StorageLink(\'loginVersion\') 绑定该值，并配合 @Watch 装饰器在值变化时触发回调。这样，用户在任意页面登录后，个人中心、消息角标、签到状态等所有相关组件都会自动刷新。')
feat('feedNeedsRefresh 列表刷新标记',
     '发帖、点赞、评论等操作后设置 feedNeedsRefresh = Date.now()，首页信息流通过 @Watch 回调自动重新拉取数据。')
feat('CurrentUser 全局单例',
     '登录后将用户信息存入全局 CurrentUser 对象，所有页面无需重复请求 /api/auth/me。')

h('6.4  全链路 JSON 安全映射', 2)
p('这是一个工程实践中容易被忽略但至关重要的问题。后端返回的 JSON 数据是纯对象（Plain Object），不包含类方法。如果前端直接在 JSON 对象上调用 record.getFormattedTime() 这样的类方法，ArkTS 运行时会抛出 TypeError。')
p('Lumina 的解决方案：在每个 API 调用的返回值处理中，执行显式的类实例映射——将 JSON 对象的每个字段逐一手动赋值给 Model 类的实例。例如 PostApi.getPosts() 返回前，将每个帖子的 JSON 数据 new PostModel() 并逐字段赋值。这一层映射虽然增加了少量样板代码，但确保了整个 UI 渲染链路中不会因为调用不存在的方法而崩溃，是大型 ArkUI 项目的关键工程实践。')

h('6.5  深色/浅色主题系统', 2)
p('Lumina 自建了完整的双主题色彩系统，不依赖 HarmonyOS 系统级主题。核心设计是一个 C 类（Color 代理），内部的每个 getter 通过 ThemeManager.isDark() 判断当前模式，返回对应颜色：')
code('class C {')
code('  static get bg(): string {')
code('    return ThemeManager.isDark() ? "#121212" : "#F6F6F6"')
code('  }')
code('  static get textPrimary(): string { ... }')
code('  // ... 16 个语义化颜色属性')
code('}')
p('组件中使用 C.bg、C.textPrimary、C.stardustGold 等语义化属性名，而非硬编码色值。用户在设置页切换主题时，ThemeManager.setMode() 更新 AppStorage 中的主题标记，所有使用了 C 类颜色属性的组件自动重绘。整个主题系统约 150 行代码，却支撑了 20 个页面的完整双主题体验。')

doc.add_page_break()

# ====================================================================
#              第七章：创新特性深度解析
# ====================================================================
h('第七章  创新特性深度解析', 1)

h('7.1  星尘经济系统', 2)
p('星尘（StarDust）是 Lumina 最具创新性的功能设计。它是一个闭环的虚拟货币激励系统，将内容质量评价权从中心化的"平台推荐算法"交还给去中心化的"社区成员投票"：')
b('获取途径：每日签到（+5）、发布帖子（+2）、被他人打赏（他人主动转移星尘给你）')
b('使用途径：打赏优质帖子（6 档金额：1/5/10/20/50/100 星尘，可附留言）')
b('等级体系：累计收到星尘达到阈值自动晋升（Lv.1 星尘新芽 → Lv.6 星海领主）')
b('经济闭环：用户产出优质内容 → 被打赏获得星尘 → 用星尘打赏他人优质内容 → 激励他人产出')

p('从产品设计角度，星尘系统解决了社区类产品的经典"冷启动问题"：新平台缺乏内容，早期用户缺乏激励。Lumina 的策略是——新用户注册即送 100 星尘，让他们有能力对早期内容进行打赏；被打赏的作者看到正反馈后更有动力继续创作；打赏者的星尘消耗完了，需要通过签到或发帖来获取更多。这个循环自然地引导用户从"内容消费者"转变为"内容生产者"。')

h('7.2  一键演示数据系统', 2)
p('/api/admin/bootstrap 端点是 Lumina 工程实践中的一大亮点。它体现了"为演示而设计"（Design for Demo）的工程思维：')
b('一键调用即可生成完整的演示场景：6 个粉丝账号（各有拟人化的昵称/简介/人格设定）+ 关注关系 + 点赞互动 + 评论互动 + 3 条官方消息全站广播 + 3 篇种子帖子')
b('利用 Hermetic（封闭式）设计：所有演示数据都在一次 API 调用中完成，不依赖外部服务')
b('可重复执行：多次调用不会产生重复数据，内置了幂等性检查')

p('更进一步，/api/admin/seed 端点可以生成 8 篇知识型种子帖子（标题涵盖《人类简史》读后感、量子计算现状、康德与黑格尔哲学、算法竞赛指南、互联网职业规划等真实学术话题），帖子内容均为中文长文，模拟真实高质量社区的初始内容状态。')

h('7.3  通知推送引擎', 2)
p('Lumina 的通知系统虽然基于轮询（而非 WebSocket 推送），但在架构设计上已经为未来的实时推送预留了扩展空间：')
b('createNotification() 是所有通知的单一入口函数——点赞、评论、打赏、官方消息均通过此函数创建通知')
b('通知类型枚举化：like | comment | reward | official 四种类型，前端按类型分类筛选')
b('自通知屏蔽：createNotification 内部自动检查 fromUserId === toUserId，不给自己发通知')
b('未读计数独立端点：前端可以低成本地高频轮询 /api/notifications/unread-count（仅返回一个整数），而只在用户主动打开消息页时才拉取完整通知列表')

doc.add_page_break()

# ====================================================================
#              第八章：部署运维与演示方案
# ====================================================================
h('第八章  部署运维与演示方案', 1)

h('8.1  开发环境部署', 2)
p('Lumina 后端的部署设计遵循"零配置启动"原则——clone 代码后仅需两步即可运行：')
b('步骤 1：cd server && npm install（安装 express 和 cors 两个依赖，约 2 秒）')
b('步骤 2：node server.js（启动服务，监听 0.0.0.0:3000，控制台打印所有可用端点列表）')
p('无需安装数据库、无需配置环境变量、无需运行初始化脚本。db.json 文件已预置了种子数据，启动即可用。')

h('8.2  模拟器联调', 2)
p('HarmonyOS 模拟器通过特殊 IP 地址 10.0.2.2 访问宿主机的 localhost。ApiClient.ets 中的 BASE_URL 默认配置为 http://10.0.2.2:3000。在 HarmonyOS 模拟器中启动应用即可直接与后端通信。')

h('8.3  真机联调', 2)
p('将 BASE_URL 改为电脑在局域网中的 IP 地址（如 http://192.168.1.100:3000），确保手机和电脑在同一 WiFi 网络下，并关闭电脑防火墙对 3000 端口的拦截。')

h('8.4  课设答辩演示指南', 2)
p('以下是推荐的演示流程，可以在 5-8 分钟内完整展示 Lumina 的全部核心功能：')
b('① 启动后端：node server.js → 展示控制台输出和 db.json 数据结构')
b('② 注册/登录：展示密码安全哈希存储和 Token 认证流程')
b('③ 首页浏览：展示信息流 + 标签筛选 + 深色/浅色主题切换')
b('④ 热榜：展示加权热度算法排序效果和金/银/铜排名徽章')
b('⑤ 发布帖子：展示图片选择、标签分类、富文本输入')
b('⑥ 互动操作：展示点赞动画、评论输入、星尘打赏面板')
b('⑦ 一键演示：调用 POST /api/admin/bootstrap → 展示通知爆炸效果')
b('⑧ 搜索功能：输入关键词，展示跨字段模糊搜索')
b('⑨ 私信聊天：展示实时对话功能')
b('⑩ 个人中心：展示等级徽章、星尘明细、社交数据')

doc.add_page_break()

# ====================================================================
#              第九章：技术成果与迭代展望
# ====================================================================
h('第九章  技术成果与迭代展望', 1)

h('9.1  技术成果统计', 2)

# 成果表格
table2 = doc.add_table(rows=12, cols=2, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
data2 = [
    ['指标', '数值'],
    ['前端页面数', '20 个独立页面 + 1 个可复用组件'],
    ['前端代码量', '约 4,500 行 ArkTS 代码'],
    ['后端 API 端点数', '25 个 RESTful 端点，分 7 个模块'],
    ['后端代码量', '约 860 行 Node.js (server.js)'],
    ['数据集合数', '8 个（users / posts / comments / starDustRecords / notifications / follows / conversations / messages）'],
    ['核心依赖数', '仅 2 个（express + cors）'],
    ['支持设备类型', 'Phone + Tablet (HarmonyOS 5.0 / API 12+)'],
    ['主题系统', '深色/浅色双主题，约 150 行自研颜色代理'],
    ['安全机制', 'PBKDF2-SHA512 加盐哈希 + Bearer Token 认证 + 数据脱敏'],
    ['激励系统', '星尘虚拟货币 + 六等级体系 + 打赏经济闭环'],
]
for r_idx, row_data in enumerate(data2):
    for c_idx, cell_text in enumerate(row_data):
        cell = table2.rows[r_idx].cells[c_idx]
        cell.text = cell_text
        for pp in cell.paragraphs:
            for run in pp.runs:
                if r_idx == 0:
                    run.bold = True
                run.font.size = Pt(10)

doc.add_paragraph()

h('9.2  技术栈全景图', 2)
p('Lumina 项目展示了一条完整的全栈技术链路：')
b('后端：Node.js → Express → RESTful API → JSON 文件数据库 → PBKDF2 密码哈希 → Bearer Token 认证')
b('前端：HarmonyOS 5.0 → ArkTS → ArkUI 声明式 UI → AppStorage 状态管理 → 组件化架构')
b('网络：HTTP JSON 通信 → 统一响应格式 → Token 自动注入 → 错误处理与用户提示')
b('工程：Git 版本控制 → 模块化项目结构 → 一键演示数据 → 零配置部署')

h('9.3  未来迭代方向', 2)
b('数据库升级：将 JSON 文件数据库迁移至 SQLite（HarmonyOS 原生支持的关系型数据库），支持事务和并发控制')
b('图片上传：实现基于 multipart/form-data 的真实图片上传功能，替代当前的纯文本图片 URL 方案')
b('实时通知：引入 WebSocket 或在 HarmonyOS 端使用 @ohos.net.webSocket 实现通知的实时推送，替代当前的前端轮询')
b('内容审核：接入基于 AI 的文本审核服务，对帖子和评论进行自动合规检测')
b('推荐算法：基于协同过滤或内容标签为用户推荐可能感兴趣的帖子')
b('性能优化：引入 Redis 缓存热榜和首页信息流，减少 JSON 文件 IO 频率')
b('鸿蒙能力集成：利用 HarmonyOS 的分布式数据管理能力实现同账号多设备无缝同步')

h('9.4  个人技术收获', 2)
p('Lumina 项目从零到一完整构建了一款知识型社区 App 的全部前后端功能。在这个过程中，我们：')
b('深入理解了 HarmonyOS ArkTS + ArkUI 的声明式开发范式，能够独立完成鸿蒙原生 App 的全流程开发')
b('掌握了 RESTful API 从需求分析 → 接口设计 → 代码实现 → 安全加固的完整流程')
b('实践了 PBKDF2 密码哈希、Token 认证、数据脱敏等后端安全技术，建立起了安全意识')
b('积累了大型 App 项目的组件化架构设计经验——20 个页面 + 状态管理 + 主题系统的工程化实践')
b('学会了如何从产品视角设计一个完整的虚拟货币激励系统——不只是写代码，更是设计行为经济学')

doc.add_paragraph()
doc.add_paragraph()

# 结语
pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('以知识为星尘，让深度思考被看见')
rr.bold = True; rr.font.size = Pt(16); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'
rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

pp2 = doc.add_paragraph()
pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr2 = pp2.add_run('— Lumina 团队 · 2026 —')
rr2.font.size = Pt(11); rr2.font.color.rgb = RGBColor(150, 150, 150)
rr2.font.name = '微软雅黑'
rr2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 保存 =====
output_path = 'Lumina_产品介绍书_v2.docx'
doc.save(output_path)
print(f'Done: {output_path}')
print(f'   共 9 章，涵盖产品定位、功能全景、后端架构（重点）、API 设计、安全体系、前端架构、创新特性、部署方案、技术成果')
