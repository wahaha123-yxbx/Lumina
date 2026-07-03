# -*- coding: utf-8 -*-
"""生成文档：从课设到工业级 —— 后端差距分析与学习路线"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os

os.chdir('C:/Lumina')
doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)

def H(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def P(doc, text, bold=False, color=None, size=10.5):
    pp = doc.add_paragraph(); r = pp.add_run(text)
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    pp.paragraph_format.space_after = Pt(4)

def B(doc, text):
    pp = doc.add_paragraph(text, style='List Bullet')
    pp.paragraph_format.space_after = Pt(2); pp.paragraph_format.line_spacing = 1.4
    for run in pp.runs:
        run.font.size = Pt(10); run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def PAGE(doc):
    doc.add_page_break()

# ===== 封面 =====
for _ in range(5): doc.add_paragraph()
t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run('从 Lumina 到工业级后端')
r.font.size = Pt(36); r.bold = True; r.font.color.rgb = RGBColor(0, 102, 255)
r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.add_paragraph()
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.add_run('课设后端 vs 大厂后端 —— 差距、原因与学习路线').font.size = Pt(14)
doc.add_paragraph(); doc.add_paragraph()
today = datetime.date.today().strftime('%Y.%m.%d')
for ml in ['作者：葛沛林', f'日期：{today}', '基于 Lumina 项目的真实复盘']:
    tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tt.add_run(ml).font.size = Pt(10.5)
PAGE(doc)

# ===== 一、引言 =====
H(doc, '一、我们做了什么，大厂做到了什么', 1)

P(doc, 'Lumina 后端是一个标准的 RESTful API 服务器：Node.js + Express，25+ 个接口，JSON 文件数据库，PBKDF2 密码哈希，自研 Token 认证。它能完美支撑一个单机、单用户、课设答辩场景。')
P(doc, '')
P(doc, '当用户量从 1 个变成 1000 万，这套架构会发生什么？', bold=True)
B(doc, 'JSON 文件每次请求全量读写 → 用户多了磁盘 IO 直接打满，请求排队等死')
B(doc, 'Token 存内存 Map → 服务重启 1000 万用户全部掉线')
B(doc, '热榜每次实时全表排序 → 100 万帖子时每次请求算几秒钟')
B(doc, '所有数据一台机器 → 这台机器挂了整个 App 不可用')
B(doc, '没有日志系统 → 出 Bug 了完全不知道哪个请求出的问题')

P(doc, '')
P(doc, '这些不是"我们做错了"——是课设场景和工业场景的天然差异。这份文档诚实地列出差距，并给出从课设走向工业级的学习路线。')

PAGE(doc)

# ===== 二、逐层对比 =====
H(doc, '二、逐层对比：Lumina vs 大厂', 1)

# ---- 2.1 ----
H(doc, '2.1  认证与安全', 2)

table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['维度', 'Lumina（当前）', '大厂标准']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)

for r, row in enumerate([
    ['密码存储', 'PBKDF2-SHA512，1000次迭代', 'bcrypt/argon2，10万+次迭代，自适应难度'],
    ['Token', '自研随机字符串，内存Map存', 'JWT+Refresh双Token，Redis集群存，支持黑名单'],
    ['传输安全', 'HTTP 明文', 'HTTPS/TLS 1.3，证书自动轮换'],
    ['暴力破解防护', '无', 'IP限流 + 验证码 + 账号锁定 + 设备指纹'],
    ['SQL注入防护', '无SQL（JSON文件）', '参数化查询 + ORM + WAF防火墙'],
    ['敏感数据', '手动解构脱敏', '字段级加密 + 审计日志 + 数据脱敏中间件'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

doc.add_paragraph()

# ---- 2.2 ----
H(doc, '2.2  数据存储', 2)

table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['维度', 'Lumina（当前）', '大厂标准']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)

for r, row in enumerate([
    ['数据库', '单个JSON文件', 'MySQL/PostgreSQL，分库分表，一主多从'],
    ['读写方式', '每次全量readDB/writeDB', 'SQL精确查询 + 索引优化 + 连接池'],
    ['事务', '不支持', 'ACID事务，BEGIN/COMMIT/ROLLBACK'],
    ['备份', '手动 cp 文件', '自动定时备份 + binlog增量 + 异地容灾'],
    ['高可用', '单机，挂了就没了', '主从切换 + 多机房部署 + 故障自动转移'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

doc.add_paragraph()

# ---- 2.3 ----
H(doc, '2.3  性能与扩展', 2)

table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['维度', 'Lumina（当前）', '大厂标准']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)

for r, row in enumerate([
    ['缓存', '无，每次实时读文件', 'Redis 缓存热数据（热榜、用户信息），命中率90%+'],
    ['热榜计算', '每次请求全量排序O(n log n)', '定时任务预计算 + Redis有序集合 ZREVRANGE'],
    ['并发处理', 'Node.js单线程处理', '多进程cluster + 容器化 + Kubernetes自动扩缩'],
    ['静态资源', '无（前端直接读本地）', 'CDN分发，用户就近访问'],
    ['消息处理', '同步直接操作', '消息队列异步解耦（Kafka/RabbitMQ）'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

doc.add_paragraph()

# ---- 2.4 ----
H(doc, '2.4  运维与监控', 2)

table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['维度', 'Lumina（当前）', '大厂标准']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)

for r, row in enumerate([
    ['部署方式', '手动 node server.js', 'Docker镜像 + K8s编排 + CI/CD自动部署'],
    ['日志', 'console.log 打印', '结构化日志 + ELK收集 + 全文检索'],
    ['监控告警', '无', 'Prometheus+Grafana看板 + 企业微信/钉钉告警'],
    ['灰度发布', '无', '滚动更新 + 金丝雀发布 + 流量切换'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

PAGE(doc)

# ===== 三、架构演进 =====
H(doc, '三、如果 Lumina 有 1000 万用户，架构会变成什么样', 1)

P(doc, '下图展示同一款 App 在三个阶段的架构变化。Lumina 当前在阶段一。', bold=True)

P(doc, '')
H(doc, '阶段一：单体应用（现在的 Lumina）', 2)
P(doc, '  一台电脑 + node server.js + 一个 JSON 文件 = 全部功能')
P(doc, '  适用：1 个开发者，1 个用户，课设答辩')

H(doc, '阶段二：前后端分离 + 数据库（毕业设计 / 创业 MVP）', 2)
P(doc, '  前端独立部署 + Node.js 集群（PM2 多进程）+ MySQL/PostgreSQL + Redis 缓存')
P(doc, '  适用：几百到几千用户，小团队开发')

H(doc, '阶段三：微服务 + 云原生（真正的大厂）', 2)
P(doc, '  网关(Nginx/Kong) → 认证服务 帖子服务 用户服务 通知服务 搜索服务（各自独立部署）')
P(doc, '  → Redis 集群 → MySQL 分库分表 + Elasticsearch 搜索引擎 → Kafka 消息队列')
P(doc, '  → Prometheus 监控 + ELK 日志 + Docker + Kubernetes 编排')

P(doc, '')
P(doc, '每一步演进解决的是同一个问题：', bold=True)
P(doc, '"上一阶段的架构在某一个维度上撑不住了"——要么是并发，要么是数据量，要么是开发效率，要么是稳定性。没有最好的架构，只有最适合当前规模的架构。')

PAGE(doc)

# ===== 四、学习路线 =====
H(doc, '四、从课设到工业级：学习路线', 1)

P(doc, '以下路线基于当前 Lumina 后端技术栈（Node.js/JavaScript 生态）延伸，按优先级排序。', bold=True)

H(doc, '第一层：夯实基础（课设做完就可以开始）', 2)
B(doc, '学 MySQL：建表、索引、JOIN、事务——SQL 是所有后端的通用语言。推荐《MySQL必知必会》')
B(doc, '学 Redis：缓存是什么、String/Hash/List/Sorted Set/ZSet 各自适用什么场景。把 Lumina 热榜改成 Redis ZSet 实现')
B(doc, '学 RESTful 深入：HTTP 状态码语义（201/204/401/403/429）、分页设计（cursor vs offset）、API 版本管理')
B(doc, '学 Docker：把你的 server.js 和 db.json 打包成镜像，一条命令在任何电脑上跑起来')

H(doc, '第二层：工程化（做一个真正能用的项目）', 2)
B(doc, '学 TypeScript：把 server.js 重写成 server.ts。类型安全 = 少一半 Bug。Express 换成 Nest.js（TS 原生框架）')
B(doc, '学 ORM：Prisma 或 TypeORM，告别手写 JSON 查询')
B(doc, '学 JWT + OAuth 2.0：理解 access token / refresh token / 第三方登录的完整流程')
B(doc, '学 Git 规范：Git Flow、Conventional Commits、Code Review 流程')
B(doc, '学 CI/CD：GitHub Actions 自动测试 + 自动部署')

H(doc, '第三层：分布式（进大厂的硬通货）', 2)
B(doc, '学消息队列：Kafka 或 RabbitMQ——异步解耦是分布式系统的核心思想')
B(doc, '学微服务：什么时候拆、怎么拆、服务间怎么通信（gRPC/HTTP/MQ）')
B(doc, '学 K8s：Pod/Service/Deployment/Ingress，理解容器编排的声明式哲学')
B(doc, '学分布式理论：CAP 定理、BASE 理论、分布式事务（Seata/Saga）')
B(doc, '学监控：Prometheus + Grafana + Loki 全链路可观测性')

H(doc, '第四层：持续精进', 2)
B(doc, '读源码：Express / Redis / Nginx 源码，理解"为什么这么设计"')
B(doc, '写博客：把学到的写出来，输出倒逼输入')
B(doc, '参与开源：从修 typo 开始，到提 feature PR')
B(doc, '做自己的产品：像 Lumina 一样，从零到一完整做完一个项目。最好的学习方式永远是"做出来"')

PAGE(doc)

# ===== 五、答辩 QA =====
H(doc, '五、答辩时的"差距问题"怎么回答', 1)

P(doc, '以下是答辩老师最可能就"差距"方向提出的问题，以及建议回答。', bold=True)

P(doc, '')
P(doc, 'Q1：你这个后端太简陋了吧？JSON 文件也算数据库？', bold=True)
P(doc, '答："这确实是一个经过权衡的工程决策。课设阶段数据量不到 1000 条，JSON 文件可以做到零配置启动、Git 版本控制友好、演示时可以直接用文本编辑器打开展示数据结构。更重要的是，我们用了依赖倒置原则——所有数据操作封装在 readDB 和 writeDB 两个函数中，未来迁移到 MySQL 只需要改这两个函数，上面 800 行业务代码一行不动。"')

P(doc, '')
P(doc, 'Q2：那你觉得和真正大厂后端的差距在哪？', bold=True)
P(doc, '答："差距在四个层面。一是数据层——我们没有真正的数据库、缓存和消息队列；二是安全层——缺少限流、HTTPS、验证码等防护；三是运维层——没有日志系统、监控告警和容器化部署；四是架构层——单体应用无法横向扩展。但我认为课设阶段把这些全做了反而是错误的——过度设计。关键是知道差距在哪、怎么补，而不是在不需要的时候硬上。"')

P(doc, '')
P(doc, 'Q3：如果让你继续做这个项目，第一个要改的是什么？', bold=True)
P(doc, '答："数据库。JSON 文件是当前最大的单点瓶颈——每次请求全量读写，数据越多越慢。我会把 db.json 迁移到 SQLite（HarmonyOS 原生支持），然后引入 Redis 缓存热榜和用户信息。这两步改完，系统能扛住至少 1000 倍的数据量增长。后面再做的话，把敏感接口加上限流，接入 HTTPS，最后用 Docker 容器化部署。"')

P(doc, '')
P(doc, 'Q4：你写这个项目最大的技术收获是什么？', bold=True)
P(doc, "答：两点。第一，我理解了'架构是演化出来的，不是设计出来的'。Lumina 从最初的几个接口慢慢长到 32 个 API + 8 个数据集合，每个新增功能都在倒逼我思考'代码怎么组织才不乱'，于是自然形成了五层架构。第二，我学到了'不做过度设计'——课设场景用 JSON 文件是明智的，用 K8s 才是愚蠢的。知道什么阶段该用什么技术，比会用多少技术更重要。")

P(doc, '')
P(doc, '')
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('— 从 Lumina 到工业级 · 完 —')
rr.font.size = Pt(13); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.save('Lumina_Industrial_Comparison.docx')
print('Done')
