# -*- coding: utf-8 -*-
"""Lumina 后端技术精要 —— 课设答辩用，直击重点难点"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os

os.chdir('C:/Lumina')
doc = Document()

# ---- 样式 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)

def H(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def P(doc, text, bold=False, color=None, size=11):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    pp.paragraph_format.space_after = Pt(4)
    pp.paragraph_format.line_spacing = 1.5

def B(doc, text):
    pp = doc.add_paragraph(text, style='List Bullet')
    pp.paragraph_format.space_after = Pt(2)
    pp.paragraph_format.line_spacing = 1.4
    for run in pp.runs:
        run.font.size = Pt(10.5); run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def CODE(doc, text):
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(1); pp.paragraph_format.line_spacing = 1.2
    r = pp.add_run('  ' + text)
    r.font.name = 'Consolas'; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(60, 60, 60)

def KEY(doc, title, content):
    """重点条目"""
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(4); pp.paragraph_format.line_spacing = 1.5
    r1 = pp.add_run('📍 ' + title + '：')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0, 102, 255)
    r1.font.name = '微软雅黑'; r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = pp.add_run(content)
    r2.font.size = Pt(10.5); r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def PAGE(doc):
    doc.add_page_break()

# ================================================================
# 封面
# ================================================================
for _ in range(5): doc.add_paragraph()
t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run('Lumina 后端技术精要')
r.font.size = Pt(42); r.bold = True; r.font.color.rgb = RGBColor(0, 102, 255)
r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.add_paragraph()
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.add_run('Node.js + Express RESTful API 服务端').font.size = Pt(14)
t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.add_run('课设答辩 · 重点难点直击').font.size = Pt(13)
doc.add_paragraph(); doc.add_paragraph()
today = datetime.date.today().strftime('%Y.%m.%d')
for ml in [f'技术栈：Node.js + Express + JSON DB    |    代码量：~860 行    |    API：25 个    |    数据集合：8 个',
           f'日期：{today}']:
    tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tt.add_run(ml).font.size = Pt(10)

PAGE(doc)

# ================================================================
# 一、后端总览
# ================================================================
H(doc, '一、后端做了什么', 1)

P(doc, '一句话概括：用 Node.js + Express 搭建了一个 RESTful API 服务器，为 Lumina App 的 20 个前端页面提供全部数据服务。服务器监听 3000 端口，HarmonyOS 端通过 HTTP 请求与后端通信。')

P(doc, '项目文件结构：', bold=True)
CODE(doc, 'server/')
CODE(doc, '├── server.js          ← 主文件，860 行，全部业务逻辑')
CODE(doc, '├── data/db.json       ← JSON 文件数据库，8 个集合')
CODE(doc, '├── package.json       ← 仅 2 个依赖：express + cors')
CODE(doc, '└── node_modules/      ← npm install 自动安装')

P(doc, '启动方式：', bold=True)
CODE(doc, 'cd server && npm install && node server.js')
CODE(doc, '# → http://localhost:3000')

P(doc, '')
KEY(doc, '核心数据流（每次请求的完整路径）',
    '客户端 ArkTS → HTTP 请求 → CORS 中间件 → JSON 解析中间件 → authMiddleware 认证中间件 → 路由匹配 → 业务处理函数(readDB→逻辑→writeDB) → 返回 JSON 响应')

PAGE(doc)

# ================================================================
# 二、技术选型
# ================================================================
H(doc, '二、技术选型（为什么这么选）', 1)

KEY(doc, 'Node.js + Express',
    '前端 ArkTS 是 TypeScript 超集，后端用 JavaScript，全栈语言统一，一个人能无缝切换。Express 是 Node.js 最成熟的 Web 框架（GitHub 60k+ stars），中间件模型直观，路由定义清晰，适合 RESTful API 快速开发。')
KEY(doc, 'JSON 文件数据库',
    '课设数据量 < 1000 条，JSON 文件完全够用。优势：零安装配置（不用装 MySQL）、Git 版本控制友好、文本编辑器直接打开查看、readDB/writeDB 两个函数封装一切。未来迁移数据库只需改这两个函数。')
KEY(doc, '仅 2 个依赖（express + cors）',
    '密码哈希用 Node.js 内置 crypto（PBKDF2），Token 自研不用 JWT 库，ID 生成自研不用 uuid。零第三方认证库 = 零供应链风险，同时展示从底层理解安全的工程能力。')

PAGE(doc)

# ================================================================
# 三、五层架构
# ================================================================
H(doc, '三、五层架构设计（重点）', 1)

P(doc, '后端代码虽然只有 860 行，但按照五层架构组织，层次清晰：')

H(doc, '第 1 层 — 中间件链', 2)
CODE(doc, 'app.use(cors())              // 允许跨域')
CODE(doc, 'app.use(express.json())      // 自动解析 JSON 请求体')
CODE(doc, 'app.use(authMiddleware)      // 从 Header 提取 Token → 解析 userId')
P(doc, 'authMiddleware 是"软认证"——解析不到 Token 不会拒绝请求，而是让 req.userId 保持 undefined，由下游业务函数自行判断是否需要登录。')

H(doc, '第 2 层 — 路由分发', 2)
P(doc, '25 个端点按 HTTP 方法 + URL 路径分发。RESTful 语义：GET=读取，POST=创建，PUT=更新。URL 使用资源名词（/api/posts、/api/users），避免动词命名。')

H(doc, '第 3 层 — 业务逻辑', 2)
P(doc, '每个路由处理函数执行：参数校验 → 读取数据 → 业务判断 → 修改数据 → 写回文件 → 返回响应。整个操作在一个函数中完成，readDB → 修改 → writeDB 形成原子操作单元。')

H(doc, '第 4 层 — 数据访问', 2)
CODE(doc, 'function readDB()  { return JSON.parse(fs.readFileSync(DB_PATH, "utf-8")) }')
CODE(doc, 'function writeDB(d) { fs.writeFileSync(DB_PATH, JSON.stringify(d, null, 2), "utf-8") }')
P(doc, '只有这两个函数封装对 db.json 的全部 IO。未来换 MySQL 只需改这两个函数——上面三层一行不用动。这就是"依赖倒置原则"。')

H(doc, '第 5 层 — 基础设施', 2)
CODE(doc, 'genId(prefix)            → 唯一 ID 生成（前缀_时间戳36进制_随机4位）')
CODE(doc, 'hashPassword(pw, salt)   → PBKDF2-SHA512 密码哈希')
CODE(doc, 'makeSalt()               → crypto.randomBytes(16) 随机盐')
CODE(doc, 'generateToken()          → Token 生成（lm_时间戳_随机16位）')
CODE(doc, 'setToken(userId)         → 存入内存 Map<token, userId>')
CODE(doc, 'calculateLevel(total)    → 六等级判定')
CODE(doc, 'createNotification(...)  → 通知创建统一入口')

PAGE(doc)

# ================================================================
# 四、安全体系（难点）
# ================================================================
H(doc, '四、安全体系（难点 & 加分项）', 1)

H(doc, '4.1  密码存储：PBKDF2-SHA512 加盐哈希', 2)
P(doc, '绝不明文存储密码。注册时对每个用户执行独立的加盐哈希处理。完整流程：', bold=True)

P(doc, '注册阶段：', bold=True)
B(doc, 'crypto.randomBytes(16) → 16 字节密码学安全随机盐（2^128 空间，碰撞概率 ≈ 0）')
B(doc, 'crypto.pbkdf2Sync(password, salt, 1000, 64, "sha512") → 64 字节哈希值')
B(doc, '存储 passwordHash 和 passwordSalt，丢弃原始密码')

P(doc, '登录阶段：', bold=True)
B(doc, '取出该用户的 salt → 对输入密码执行相同哈希 → 比对 hash === storedHash')
B(doc, '匹配成功 → 生成 Token → 返回脱敏用户信息')
B(doc, '匹配失败 → 返回"密码错误"')

P(doc, '')
KEY(doc, '安全核心',
    '即使两个用户设置完全相同的密码，由于各自拥有独立的随机盐值，数据库中的哈希值完全不同。攻击者拿到 db.json 也无法通过彩虹表反推密码——必须对每个用户的盐值单独暴力枚举，计算上不可行。')
KEY(doc, '为什么是 1000 次迭代',
    '每次迭代增加攻击者暴力破解的成本。PBKDF2 的设计目标就是"慢"——对正常登录的一次哈希几乎无感知（<1ms），但对暴力枚举上亿次尝试则成本极高。生产环境推荐 10 万次以上，课设用 1000 次作为演示。')

H(doc, '4.2  Token 认证：自研方案，不用 JWT', 2)
P(doc, '自研 Token 而非 JWT，理由有三：')
B(doc, 'Token 不含任何用户信息——格式为 lm_{时间戳}_{随机16位}，纯随机字符串，即使泄露也无法解码')
B(doc, '可随时撤销——服务端内存 Map<token, userId>，删掉 key 即可强制下线。JWT 无状态，签发后无法撤销')
B(doc, '单点登录——setToken() 生成新 Token 前自动清除该用户所有旧 Token')
CODE(doc, '// Token 存储与生成')
CODE(doc, 'const tokenStore = new Map()  // token → userId')
CODE(doc, "function generateToken() { return 'lm_' + Date.now().toString(36) + '_' + Math.random().toString(36).substr(2, 16) }")

H(doc, '4.3  数据脱敏', 2)
P(doc, '所有返回用户的接口都用解构赋值排除敏感字段：')
CODE(doc, 'const { passwordHash, passwordSalt, ...safeUser } = user')
CODE(doc, 'res.json({ code: 0, data: safeUser })')
P(doc, '一行代码确保密码哈希和盐值绝对不会泄露到客户端。')

H(doc, '4.4  6 道业务安全校验', 2)
B(doc, '密码长度校验：password.length < 6 → 拒绝注册')
B(doc, '学号唯一性：注册时检查是否已存在 → 拒绝重复注册')
B(doc, '自关注防护：userId === targetId → 拒绝关注')
B(doc, '重复关注防护：关注关系已存在 → 拒绝重复关注')
B(doc, '余额校验：打赏时 starDustBalance < amount → 拒绝交易')
B(doc, '自通知屏蔽：fromUserId === toUserId → 跳过通知创建（不给自己发通知）')

PAGE(doc)

# ================================================================
# 五、数据模型
# ================================================================
H(doc, '五、数据模型设计', 1)

P(doc, 'JSON 数据库包含 8 个集合，对应关系型数据库的 8 张表。设计原则：高频字段适当冗余，避免"多表关联"查询。')

# 简表
table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['集合', '核心字段', '设计要点']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)

data = [
    ['users', 'id, studentId, nickname, passwordHash, passwordSalt,\nstarDustBalance, totalStarDustReceived, level(1-6),\npostCount, followingCount, followerCount',
     '社交计数冗余存储，避免每次\n展示用户卡片都 COUNT 关联表'],
    ['posts', 'id, title, content, images[], authorId, authorName,\ntag(10种分类), likeCount, commentCount,\nstarDustCount, viewCount, isTop, likedBy[]',
     'authorName 冗余避免 JOIN；\nlikedBy[] 替代独立点赞表'],
    ['comments', 'id, postId, content, authorId, authorName,\ncreateTime, likeCount, likedBy[]', '通过 postId 关联帖子'],
    ['starDust\nRecords', 'id, fromUserId, toUserId, postId, amount,\nmessage, createTime', '星尘经济系统的"账本"，\n记录每笔流转'],
    ['notifications', 'id, type(like|comment|reward|official),\nfromUserId, toUserId, postId, isRead', 'type 字段区分 4 种通知类型'],
    ['follows', 'fromUserId, toUserId, createTime', '极简二元组，双向查询'],
    ['conversations', 'id, user1, user2, lastMessage, lastTime', 'lastMessage 冗余最后一条消息预览'],
    ['messages', 'id, conversationId, fromUserId, toUserId,\ncontent, createTime, isRead', '通过 conversationId 关联会话'],
]
for r, row in enumerate(data):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(8.5)

P(doc, '')
KEY(doc, '关键设计决策——冗余字段',
    'postCount/followingCount/followerCount 内嵌在 users 表中，而非每次 COUNT。原因：用户卡片是最高频展示场景，每次都 COUNT 所有帖子/关注是性能灾难。"空间换时间"是后端优化的经典策略。')
KEY(doc, 'likedBy 数组设计',
    '用 likedBy: ["u_001","u_002",...] 替代独立的点赞关系表。点赞 push，取消点赞 splice。O(n) 查找在课设数据量下完全可接受，且代码简洁。')

PAGE(doc)

# ================================================================
# 六、API 接口速览
# ================================================================
H(doc, '六、25 个 API 端点速览', 1)

P(doc, '所有接口统一规范：URL 以 /api/ 开头，请求体 JSON 格式，响应体 { code: 0, data: ..., message: ... }。code=0 成功，code=1 业务错误。')

H(doc, '认证模块（3 个）', 2)
CODE(doc, 'POST /api/auth/register    注册（学号+密码→PBKDF2哈希→Token）')
CODE(doc, 'POST /api/auth/login       登录（验证密码→Token）')
CODE(doc, 'GET  /api/auth/me          当前用户信息（需 Token）')

H(doc, '用户模块（10 个）', 2)
CODE(doc, 'GET    /api/users/:id              用户详情')
CODE(doc, 'PUT    /api/users/:id              更新资料（部分更新）')
CODE(doc, 'POST   /api/users/:id/checkin      每日签到 +5 星尘')
CODE(doc, 'POST   /api/users/:id/follow       关注  |  unfollow  取消关注')
CODE(doc, 'GET    /api/users/:id/follow-status 关注状态查询')
CODE(doc, 'GET    /api/users/:id/followers    粉丝列表  |  following  关注列表')
CODE(doc, 'GET    /api/users/:id/posts        用户帖子  |  rewards  星尘记录')

H(doc, '帖子模块（8 个）', 2)
CODE(doc, 'GET    /api/posts              帖子列表（?tag=分类筛选，置顶优先）')
CODE(doc, 'POST   /api/posts              发布帖子（+2 星尘奖励）')
CODE(doc, 'GET    /api/posts/hot          热榜（加权热度实时排序）')
CODE(doc, 'GET    /api/posts/:id          帖子详情（浏览+1）')
CODE(doc, 'POST   /api/posts/:id/like     点赞/取消 Toggle + 通知')
CODE(doc, 'GET    /api/posts/:id/comments 评论列表')
CODE(doc, 'POST   /api/posts/:id/comments 发表评论 + 通知')
CODE(doc, 'POST   /api/posts/:id/reward   星尘打赏（余额校验+转账+升级+通知）')

H(doc, '通知 + 消息（6 个）', 2)
CODE(doc, 'GET    /api/notifications            通知列表（?userId=）')
CODE(doc, 'GET    /api/notifications/unread-count 未读计数')
CODE(doc, 'POST   /api/notifications/read       标记已读')
CODE(doc, 'POST   /api/messages/send            发送私信')
CODE(doc, 'GET    /api/messages/conversations   会话列表')
CODE(doc, 'GET    /api/messages/:convId         会话消息')

H(doc, '搜索 + 管理（4 个）', 2)
CODE(doc, 'GET    /api/search?q=关键字         全字段模糊搜索')
CODE(doc, 'POST   /api/admin/seed              生成 8 篇知识型种子帖子')
CODE(doc, 'POST   /api/admin/bootstrap         一键生成演示数据')
CODE(doc, 'POST   /api/admin/broadcast         全站广播官方消息')

PAGE(doc)

# ================================================================
# 七、核心难点拆解
# ================================================================
H(doc, '七、三个核心技术难点', 1)

H(doc, '难点 1：星尘打赏的 7 步事务操作', 2)
P(doc, '打赏是业务逻辑最复杂的端点，涉及 7 步操作且必须全部成功：')
CODE(doc, 'POST /api/posts/:id/reward  { fromUserId, amount, message }')
CODE(doc, '')
CODE(doc, '① 查帖子：db.posts.find(p => p.id === postId)')
CODE(doc, '② 查打赏者：db.users.find(u => u.id === fromUserId)')
CODE(doc, '③ 余额校验：if (fromUser.starDustBalance < amount) → 400')
CODE(doc, '④ 扣款：fromUser.starDustBalance -= amount')
CODE(doc, '⑤ 入账+升级：toUser.totalStarDustReceived += amount')
CODE(doc, '   → calculateLevel(total) → 更新 level 和 levelTitle')
CODE(doc, '⑥ 记录流水：db.starDustRecords.unshift(record)')
CODE(doc, '⑦ 创建通知：createNotification(db, "reward", ...)')
CODE(doc, '⑧ 持久化：writeDB(db)')

P(doc, '')
KEY(doc, '当前方案的局限与改进方向',
    'JSON 文件不支持事务回滚——如果第 ⑦ 步失败，前 6 步已修改的数据无法回滚。课设阶段单用户操作几乎不会触发此问题，但在答辩中主动指出这一点并说明改进方案（迁移 SQLite 使用 BEGIN/COMMIT/ROLLBACK 事务），展示工程意识。')

H(doc, '难点 2：热榜加权算法', 2)
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('热度 = 点赞×2 + 评论×3 + 星尘×5 + 浏览×0.1')
rr.bold = True; rr.font.size = Pt(13); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
P(doc, '')
B(doc, '星尘(5x) 权重最高：打赏者付出了真实虚拟资产，是最强的内容认可信号')
B(doc, '评论(3x) 次之：评论需要用户主动输入文字，参与度高于简单点赞')
B(doc, '点赞(2x)：门槛最低的互动行为')
B(doc, '浏览(0.1x) 权重极低：防止"标题党"靠点击量刷榜——没人互动的高点击内容不会排到前面')
P(doc, '')
CODE(doc, '// 排序实现（server.js 第 184-189 行）')
CODE(doc, 'posts.sort((a, b) => {')
CODE(doc, '  const hotA = a.likeCount*2 + a.commentCount*3 + a.starDustCount*5 + a.viewCount*0.1')
CODE(doc, '  const hotB = b.likeCount*2 + b.commentCount*3 + b.starDustCount*5 + b.viewCount*0.1')
CODE(doc, '  return hotB - hotA  // 降序')
CODE(doc, '})')

H(doc, '难点 3：通知引擎的统一入口设计', 2)
P(doc, 'createNotification() 是所有通知的单一入口函数，被点赞/评论/打赏/广播 4 个场景复用：')
CODE(doc, 'function createNotification(db, type, fromUserId, fromUserName, toUserId, postId, postTitle) {')
CODE(doc, '  if (fromUserId === toUserId) return  // 关键：不给自己发通知')
CODE(doc, '  if (!db.notifications) db.notifications = []')
CODE(doc, '  db.notifications.unshift({ id: genId("n"), type, fromUserId, ... })')
CODE(doc, '}')
P(doc, '这个设计的好处：所有通知逻辑集中管理，新增通知类型只需在一处修改；自通知屏蔽逻辑只写一次，不会遗漏。')

PAGE(doc)

# ================================================================
# 八、亮点功能
# ================================================================
H(doc, '八、课设亮点', 1)

KEY(doc, '星尘经济闭环',
    '注册送 100 星尘 → 签到每日 +5 → 发帖 +2 → 被打赏获得星尘 → 用星尘打赏他人 → 被打赏者升级 → 社区正向循环。这是一个完整的行为经济学设计。')
KEY(doc, '一键演示系统',
    '/api/admin/bootstrap 一条命令生成：6 个拟人化粉丝 + 全部关注 + 36+ 条互动 + 3 条官方消息。答辩时 1 秒展示"消息爆炸"效果，无需手动造数据。')
KEY(doc, '自研安全体系',
    '零第三方认证库：PBKDF2 用 Node.js 内置 crypto，Token 自研不用 JWT。展示的是"理解底层原理"而非"会用 npm 装库"。')
KEY(doc, '五层架构 + 依赖倒置',
    '860 行代码按五层架构组织，数据访问层只有 readDB/writeDB 两个函数。未来换数据库只需改这两个函数，上面三层不受影响。')
KEY(doc, '统一响应格式',
    '全部 25 个端点返回 { code, data, message }。前端只需检查 code 字段，0=成功 1=失败。这是工业界的标准实践。')

PAGE(doc)

# ================================================================
# 九、答辩 QA
# ================================================================
H(doc, '九、答辩高频问题预案', 1)

KEY(doc, 'Q1：为什么不用 MySQL？',
    '课设数据量小（<1000条）、单用户操作、无并发——JSON 文件方案完美匹配。同时 readDB/writeDB 封装了全部 IO，未来迁移只需改这两个函数。这是务实的工程决策，不是技术缺陷。')
KEY(doc, 'Q2：安全性如何保证？',
    '三层防护：① PBKDF2-SHA512 加盐哈希（每个用户独立 16 字节随机盐）② Bearer Token 自研方案（Token 纯随机，不含用户信息，可随时撤销）③ 数据脱敏（解构赋值排除密码字段）。外加 6 道业务校验。')
KEY(doc, 'Q3：10 万用户时哪里先崩？',
    '三个瓶颈：① JSON 全量读写 → 换 SQLite/MySQL ② 热榜实时排序 O(n log n) → Redis 缓存 + 定时任务 ③ likedBy[] 的 O(n) 查找 → 换 Set 或独立点赞表。我清楚这些局限，是刻意在课设范围内做的取舍。')
KEY(doc, 'Q4：通知为什么用轮询而不用 WebSocket？',
    '当前用独立轻量端点 /api/notifications/unread-count（只返回一个整数）做轮询，开销极低。课设阶段够用。HarmonyOS 端有 @ohos.net.webSocket API，后端加 socket.io 即可升级为实时推送。')
KEY(doc, 'Q5：打赏操作如果中途失败怎么办？',
    '当前 JSON 文件不支持事务回滚。改进方案：迁移到 SQLite 后用 BEGIN TRANSACTION / COMMIT / ROLLBACK 保证原子性。答辩时主动指出现有局限和改进方案，比隐瞒问题更体现工程素养。')
KEY(doc, 'Q6：你写的最满意的代码是哪个？',
    '两个：createNotification 的统一入口设计——4 种通知类型全走一个函数，自屏蔽逻辑只写一次；bootstrap 一键演示系统——1 次 API 调用的背后是创建 6 个用户 + 6 个关注 + 36+ 条互动 + 3 条广播，体现了"为演示而设计"的工程思维。')

P(doc, '')
P(doc, '')
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('— END —')
rr.font.size = Pt(14); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.save('Lumina_后端技术精要.docx')
print('Done: 后端技术精要.docx')
