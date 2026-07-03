# -*- coding: utf-8 -*-
"""生成 Lumina 后端管理手册 —— 开发调试 & 答辩运维指南"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os

os.chdir('C:/Lumina')
doc = Document()

# 样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.4
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

def H(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def P(doc, text, bold=False, color=None, size=10.5):
    pp = doc.add_paragraph(); r = pp.add_run(text)
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    pp.paragraph_format.space_after = Pt(3)

def CMD(doc, text):
    """命令块"""
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(2); pp.paragraph_format.line_spacing = 1.2
    r = pp.add_run('  $ ' + text)
    r.font.name = 'Consolas'; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(30, 30, 30)

def NOTE(doc, text):
    pp = doc.add_paragraph(); r = pp.add_run('💡 ' + text)
    r.font.size = Pt(9.5); r.font.name = '微软雅黑'
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pp.paragraph_format.space_after = Pt(3)

def PAGE(doc):
    doc.add_page_break()

# ===== 封面 =====
for _ in range(5): doc.add_paragraph()
t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t1.add_run('Lumina 后端运维手册')
r.font.size = Pt(38); r.bold = True; r.font.color.rgb = RGBColor(0, 102, 255)
r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.add_paragraph()
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.add_run('Node.js + Express RESTful API Server').font.size = Pt(13)
doc.add_paragraph()
today = datetime.date.today().strftime('%Y.%m.%d')
for ml in ['启动 / 调试 / 测试 / 备份 / 答辩演示', f'版本 1.0  |  {today}']:
    tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tt.add_run(ml).font.size = Pt(10)
PAGE(doc)

# ===== 一、项目结构 =====
H(doc, '一、项目文件结构', 1)
P(doc, '后端共3个核心文件 + 1个依赖目录，结构极简：')
CMD(doc, 'server/')
CMD(doc, '├── server.js          ← 主程序（约800行），全部API + 业务逻辑')
CMD(doc, '├── data/db.json       ← JSON文件数据库（8个数据集合）')
CMD(doc, '├── package.json       ← 依赖声明（仅express + cors两个）')
CMD(doc, '└── node_modules/      ← npm install 自动生成（不要手动改）')
P(doc, '')
NOTE(doc, 'db.json 是数据核心，建议定期备份。node_modules 可从 git 仓库忽略，clone 后 npm install 即可重建。')

PAGE(doc)

# ===== 二、启动与停止 =====
H(doc, '二、启动 & 停止', 1)

H(doc, '2.1  首次启动（安装依赖）', 2)
CMD(doc, 'cd C:/Lumina/server        # 进入后端目录')
CMD(doc, 'npm install                # 安装依赖（仅需一次）')
CMD(doc, 'node server.js             # 启动服务')
NOTE(doc, '之后每次启动只需最后一条命令，无需重复 npm install。')

H(doc, '2.2  日常启动', 2)
CMD(doc, 'cd C:/Lumina/server')
CMD(doc, 'node server.js')
P(doc, '')
P(doc, '启动成功标志（控制台输出）：', bold=True)
CMD(doc, '✨ Lumina API Server running at http://localhost:3000')
CMD(doc, '   POST   /api/auth/register    - 注册（学号+密码）')
CMD(doc, '   POST   /api/auth/login       - 登录')
CMD(doc, '   ...（后面列出所有可用端点）')

H(doc, '2.3  停止服务', 2)
CMD(doc, '在跑 node server.js 的终端窗口按  Ctrl + C')
P(doc, '')
NOTE(doc, '改了 server.js 代码后必须重启才能生效：Ctrl+C → node server.js（改 db.json 不需要重启）。')

H(doc, '2.4  终端选择', 2)
P(doc, 'node server.js 在 Git Bash、PowerShell、cmd 中均可运行，效果完全一样。推荐用 Git Bash，因为后续 curl 测试命令在 Git Bash 中直接可用。')

PAGE(doc)

# ===== 三、日常测试命令 =====
H(doc, '三、API 测试命令（GET 查看类）', 1)

P(doc, '以下命令均在第二个终端窗口运行（第一个窗口保持 node server.js 运行）。', bold=True)

H(doc, '3.1  帖子相关', 2)
CMD(doc, 'curl http://localhost:3000/api/posts                          # 全部帖子')
CMD(doc, 'curl "http://localhost:3000/api/posts?tag=课程交流"            # 按标签筛选')
CMD(doc, 'curl http://localhost:3000/api/posts/hot                      # 热榜')
CMD(doc, 'curl http://localhost:3000/api/posts/p_001                    # 帖子详情(p_001换成实际ID)')
CMD(doc, 'curl http://localhost:3000/api/posts/p_001/comments           # 帖子评论')

H(doc, '3.2  用户相关', 2)
CMD(doc, 'curl http://localhost:3000/api/users/u_001                    # 用户信息')
CMD(doc, 'curl http://localhost:3000/api/users/u_001/posts              # 用户帖子')
CMD(doc, 'curl http://localhost:3000/api/users/u_001/rewards            # 星尘记录')
CMD(doc, 'curl http://localhost:3000/api/users/u_001/followers          # 粉丝列表')
CMD(doc, 'curl http://localhost:3000/api/users/u_001/following          # 关注列表')
CMD(doc, 'curl "http://localhost:3000/api/users/u_001/follow-status?userId=u_002"   # 关注状态')

H(doc, '3.3  通知 & 私信', 2)
CMD(doc, 'curl "http://localhost:3000/api/notifications?userId=u_001"    # 通知列表')
CMD(doc, 'curl "http://localhost:3000/api/notifications/unread-count?userId=u_001"  # 未读数')
CMD(doc, 'curl "http://localhost:3000/api/messages/conversations?userId=u_001"       # 会话列表')

H(doc, '3.4  搜索', 2)
CMD(doc, 'curl "http://localhost:3000/api/search?q=算法"                  # 全字段模糊搜索')

H(doc, '3.5  返回结果解读', 2)
P(doc, '所有接口返回统一 JSON 格式：')
CMD(doc, '{ "code": 0,          # 0=成功，1=失败')
CMD(doc, '  "data": { ... },     # 数据内容')
CMD(doc, '  "message": "..." }   # 提示信息')

PAGE(doc)

# ===== 四、POST/PUT 测试命令 =====
H(doc, '四、API 测试命令（POST/PUT 操作类）', 1)

H(doc, '4.1  认证', 2)
CMD(doc, '# 注册新用户')
CMD(doc, 'curl -X POST http://localhost:3000/api/auth/register \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"studentId":"2024001","password":"123456","nickname":"测试用户"}\'')
CMD(doc, '')
CMD(doc, '# 登录')
CMD(doc, 'curl -X POST http://localhost:3000/api/auth/login \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"studentId":"123456","password":"123456"}\'')

H(doc, '4.2  帖子操作', 2)
CMD(doc, '# 发帖（authorId 用登录返回的 user.id）')
CMD(doc, 'curl -X POST http://localhost:3000/api/posts \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"title":"测试标题","content":"正文内容","authorId":"u_001","tag":"课程交流"}\'')
CMD(doc, '')
CMD(doc, '# 点赞（再次调用自动取消点赞）')
CMD(doc, 'curl -X POST http://localhost:3000/api/posts/p_001/like \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"userId":"u_001"}\'')
CMD(doc, '')
CMD(doc, '# 发评论')
CMD(doc, 'curl -X POST http://localhost:3000/api/posts/p_001/comments \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"content":"写得太好了！","authorId":"u_001"}\'')
CMD(doc, '')
CMD(doc, '# 打赏')
CMD(doc, 'curl -X POST http://localhost:3000/api/posts/p_001/reward \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"fromUserId":"u_001","amount":10,"message":"好文章！"}\'')

H(doc, '4.3  用户操作', 2)
CMD(doc, '# 签到')
CMD(doc, 'curl -X POST http://localhost:3000/api/users/u_001/checkin')
CMD(doc, '')
CMD(doc, '# 关注')
CMD(doc, 'curl -X POST http://localhost:3000/api/users/u_003/follow \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"userId":"u_001"}\'')
CMD(doc, '')
CMD(doc, '# 取消关注')
CMD(doc, 'curl -X POST http://localhost:3000/api/users/u_003/unfollow \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"userId":"u_001"}\'')
CMD(doc, '')
CMD(doc, '# 编辑个人资料')
CMD(doc, 'curl -X PUT http://localhost:3000/api/users/u_001 \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"nickname":"新昵称","bio":"新简介","campus":"星光大学"}\'')
CMD(doc, '')
CMD(doc, '# 标记通知已读（单条）')
CMD(doc, 'curl -X POST http://localhost:3000/api/notifications/read \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"userId":"u_001","notificationId":"n_001"}\'')
CMD(doc, '')
CMD(doc, '# 标记全部已读')
CMD(doc, 'curl -X POST http://localhost:3000/api/notifications/read \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"userId":"u_001"}\'')

PAGE(doc)

# ===== 五、答辩演示命令 =====
H(doc, '五、答辩演示专用命令', 1)

H(doc, '5.1  一键生成演示数据（最重要）', 2)
CMD(doc, 'curl -X POST http://localhost:3000/api/admin/bootstrap \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"targetStudentId":"123456"}\'')
P(doc, '')
P(doc, '这条命令自动完成以下全部操作：', bold=True)
P(doc, '  ① 创建 6 个拟人化粉丝账号（算法爱好者、哲学系小张、量子萌新、读书人小王、职场观察者、思辨者）')
P(doc, '  ② 所有粉丝关注目标用户 → 粉丝数从 0 变成 6')
P(doc, '  ③ 粉丝对目标用户的帖子逐一点赞 + 评论（6 种话术轮换）')
P(doc, '  ④ 生成 36+ 条通知（点赞通知 + 评论通知）')
P(doc, '  ⑤ 全站推送 3 条官方消息')
P(doc, '')
NOTE(doc, '前提：目标学号必须已经注册。先在 App 中注册，再执行此命令。')

H(doc, '5.2  生成种子帖子', 2)
CMD(doc, 'curl -X POST http://localhost:3000/api/admin/seed')
P(doc, '生成 8 篇知识型帖子（《人类简史》读后感、量子计算现状、康德哲学比较、算法竞赛指南等），标题和正文均为真实学术内容。')
NOTE(doc, '需要在 db.json 中存在学号 123456 的用户。如果没有，改为 db.users[0]（任意已注册用户）。')

H(doc, '5.3  全站广播官方消息', 2)
CMD(doc, 'curl -X POST http://localhost:3000/api/admin/broadcast \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"title":"Lumina社区公告","content":"欢迎同学们参与讨论！"}\'')
P(doc, '遍历所有用户，每人创建一条 official 类型通知。')

H(doc, '5.4  批量生成粉丝', 2)
CMD(doc, 'curl -X POST http://localhost:3000/api/admin/generate-followers \\')
CMD(doc, '  -H "Content-Type: application/json" \\')
CMD(doc, '  -d \'{"targetUserId":"u_001","count":20}\'')
P(doc, '从已有用户中随机选取指定数量关注目标用户。')

H(doc, '5.5  推荐答辩演示顺序', 2)
P(doc, '① 确保后端已启动：node server.js')
P(doc, '② App 中注册账号（学号 123456）')
P(doc, '③ 执行 bootstrap 命令（生成 6 粉丝 + 36 互动 + 官方消息）')
P(doc, '④ 执行 seed 命令（生成 8 篇知识型帖子）')
P(doc, '⑤ 打开 App → 首页信息流有大量帖子 → 热榜有排名 → 消息页 36+ 未读')
P(doc, '⑥ 演示：浏览帖子 / 点赞 / 打赏 / 评论 / 搜索 / 私信 / 个人中心')

PAGE(doc)

# ===== 六、数据备份 =====
H(doc, '六、数据备份 & 恢复', 1)

H(doc, '6.1  保存快照', 2)
CMD(doc, '# 带日期命名，方便识别')
CMD(doc, 'cp C:/Lumina/server/data/db.json C:/Lumina/server/data/db_backup_20250625.json')
P(doc, '')
NOTE(doc, '建议答辩前、执行 bootstrap 前、大改数据前各备份一次。')

H(doc, '6.2  恢复数据', 2)
CMD(doc, 'cp C:/Lumina/server/data/db_backup_20250625.json C:/Lumina/server/data/db.json')
P(doc, '覆盖后无需重启后端，下次 API 请求自动读取恢复后的数据。')

H(doc, '6.3  直接查看数据', 2)
P(doc, '用任意文本编辑器（记事本/VSCode）打开 server/data/db.json，所有用户、帖子、评论、通知都在里面，JSON 格式直观可读。')

PAGE(doc)

# ===== 七、模拟器 & 真机切换 =====
H(doc, '七、模拟器 & 真机联调', 1)

H(doc, '7.1  模拟器', 2)
P(doc, '前端文件：entry/src/main/ets/common/ApiClient.ets（第19行）')
CMD(doc, 'static readonly BASE_URL: string = \'http://10.0.2.2:3000\'')
NOTE(doc, '10.0.2.2 是 Android 模拟器映射宿主机的固定 IP，无需修改。')

H(doc, '7.2  真机', 2)
P(doc, '将 BASE_URL 改为电脑局域网 IP：')
CMD(doc, 'static readonly BASE_URL: string = \'http://192.168.1.100:3000\'')
P(doc, '')
P(doc, '查 IP 方法：', bold=True)
CMD(doc, 'ipconfig | grep -i ipv4         # Git Bash')
CMD(doc, 'ipconfig                       # cmd/PowerShell（找 IPv4 地址）')
P(doc, '')
P(doc, '条件：', bold=True)
P(doc, '  • 手机和电脑连接同一个 WiFi')
P(doc, '  • 后端以 node server.js 保持运行')
P(doc, '  • Windows 防火墙放行 3000 端口（或临时关闭防火墙）')
P(doc, '  • 改完 BASE_URL 后需重新编译 App')

PAGE(doc)

# ===== 八、常见问题排查 =====
H(doc, '八、常见问题排查', 1)

H(doc, 'Q1: 启动报错 "Cannot find module"', 2)
CMD(doc, 'cd C:/Lumina/server && npm install')
P(doc, '依赖没装或 node_modules 被误删，重新安装即可。')

H(doc, 'Q2: 启动报错 "port 3000 already in use"', 2)
CMD(doc, '# 查找占用 3000 端口的进程并杀掉')
CMD(doc, 'netstat -ano | findstr :3000')
CMD(doc, 'taskkill /PID 进程号 /F')
P(doc, '说明之前的 node 进程没关干净。')

H(doc, 'Q3: App 首页空白 / 数据加载不出来', 2)
P(doc, '  ① 确认 node server.js 在运行（终端里有输出才算）')
P(doc, '  ② 浏览器访问 http://localhost:3000/api/posts 验证后端是否正常')
P(doc, '  ③ 检查 ApiClient.ets 中的 BASE_URL 是否正确（模拟器=10.0.2.2，真机=局域网IP）')
P(doc, '  ④ 检查手机/模拟器是否能 ping 通后端地址')

H(doc, 'Q4: 改代码后不生效', 2)
CMD(doc, 'Ctrl+C 停掉 → 重新 node server.js')
P(doc, '改了 server.js 必须重启。改了 db.json 不需要重启。')

H(doc, 'Q5: 浏览器直接访问 POST 接口报错', 2)
P(doc, '浏览器地址栏只能发 GET 请求。POST/PUT 接口需用 curl 或 Postman 或浏览器 F12 控制台的 fetch。')

H(doc, 'Q6: 数据玩坏了想重置', 2)
CMD(doc, 'cp C:/Lumina/server/data/db_backup.json C:/Lumina/server/data/db.json')
P(doc, '用之前备份的干净版本覆盖即可。')

P(doc, '')
P(doc, '')
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('— Lumina 后端运维手册 · 完 —')
rr.font.size = Pt(12); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.save('Lumina_Backend_Manual.docx')
print('Done')
