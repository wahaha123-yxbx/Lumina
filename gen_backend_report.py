# -*- coding: utf-8 -*-
"""
生成两份报告：
  1. Lumina后端技术深度解读.docx  —— 逐层拆解后端，答辩用
  2. Lumina小组分工报告.docx      —— 任务分配说明
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os

os.chdir('C:/Lumina')

# ============================================================
#                       工具函数
# ============================================================
def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.6
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    return doc

def H(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def P(doc, text, bold=False, color=None, size=11, align=None):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    pp.paragraph_format.space_after = Pt(6)
    pp.paragraph_format.line_spacing = 1.6
    if align: pp.alignment = align
    return pp

def B(doc, text, indent=0):
    pp = doc.add_paragraph(text, style='List Bullet')
    pp.paragraph_format.space_after = Pt(3)
    pp.paragraph_format.line_spacing = 1.5
    for run in pp.runs:
        run.font.size = Pt(10.5); run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if indent > 0:
        pp.paragraph_format.left_indent = Cm(1.2 * indent)
    return pp

def CODE(doc, text):
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(1)
    pp.paragraph_format.line_spacing = 1.2
    r = pp.add_run('  ' + text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(60, 60, 60)

def FEAT(doc, name, desc):
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(4); pp.paragraph_format.line_spacing = 1.5
    r1 = pp.add_run('▸ ' + name + '：')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0, 102, 255)
    r1.font.name = '微软雅黑'; r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = pp.add_run(desc)
    r2.font.size = Pt(10.5); r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def BOX(doc, title, lines):
    P(doc, '▌ ' + title, bold=True, color=(0, 102, 255), size=12)
    for line in lines: B(doc, line)

def PAGE(doc):
    doc.add_page_break()

def cover(doc, title, subtitle, desc_lines):
    for _ in range(4): doc.add_paragraph()
    t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run(title)
    r.font.size = Pt(48); r.bold = True; r.font.color.rgb = RGBColor(0, 102, 255)
    r.font.name = 'Georgia'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    t1b = doc.add_paragraph(); t1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1b = t1b.add_run(subtitle)
    r1b.font.size = Pt(16); r1b.font.color.rgb = RGBColor(100, 100, 100)
    r1b.font.name = '微软雅黑'; r1b.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph(); doc.add_paragraph()
    today = datetime.date.today().strftime('%Y 年 %m 月 %d 日')
    for ml in desc_lines + [f'日期：{today}', '技术栈：Node.js + Express + JSON DB']:
        tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = tt.add_run(ml); rr.font.size = Pt(10.5)
        rr.font.color.rgb = RGBColor(130, 130, 130)
        rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    PAGE(doc)


# ============================================================
#           报告一：后端技术深度解读
# ============================================================
doc = setup_doc()
cover(doc, 'Lumina 后端技术', '深度解读报告',
      ['—— 面向课设答辩的完整技术讲解',
       '涵盖架构设计、API 实现、安全体系、数据模型、部署方案'])

# ===== 目录 =====
H(doc, '目  录', 1)
toc = [
    '一、后端项目全景概览',
    '二、技术选型深度论证',
    '三、服务器架构逐层拆解',
    '四、数据模型完整设计',
    '五、认证与安全体系（答辩重点）',
    '六、API 端点逐一详解（25 个接口）',
    '七、核心业务逻辑剖析',
    '八、管理后台与演示数据系统',
    '九、部署与运维指南',
    '十、答辩 QA 预判',
]
for t in toc: P(doc, t)
PAGE(doc)

# ====================================================================
# 一、后端项目全景概览
# ====================================================================
H(doc, '一、后端项目全景概览', 1)

P(doc, 'Lumina 后端是一个基于 Node.js + Express 框架构建的 RESTful API 服务器。整个后端由 3 个文件组成，代码总量约 860 行，却完整支撑了 20 个前端页面的全部数据需求。', bold=True)

BOX(doc, '项目文件结构', [
    'server/'
    '  ├── server.js        ← 主服务文件（约 860 行），包含全部业务逻辑'
    '  ├── data/db.json     ← JSON 文件数据库（8 个数据集合）'
    '  ├── package.json     ← 依赖声明（仅 express + cors 两个依赖）'
    '  └── node_modules/    ← npm 自动安装的依赖包',
])

P(doc, '核心数据流（每一次请求的完整路径）：', bold=True)
B(doc, '① 客户端（HarmonyOS App）通过 @ohos.net.http 发起 HTTP 请求')
B(doc, '② 请求首先经过 CORS 中间件（允许跨域）')
B(doc, '③ 然后经过 JSON 解析中间件（req.body 自动解析）')
B(doc, '④ 然后经过认证中间件（从 Header 提取 Token → 解析 userId）')
B(doc, '⑤ 路由匹配到对应的处理函数（app.get/app.post 定义的路由）')
B(doc, '⑥ 处理函数中：readDB() 读取数据 → 业务逻辑 → writeDB() 写回')
B(doc, '⑦ 返回统一格式 JSON 响应：{ code: 0, data: ..., message: ... }')
B(doc, '⑧ 客户端 ApiClient 解析响应 → 类型安全映射 → 渲染 UI')

PAGE(doc)

# ====================================================================
# 二、技术选型深度论证
# ====================================================================
H(doc, '二、技术选型深度论证', 1)

P(doc, '这部分是答辩时必须讲清楚的——为什么选择这些技术，而不是其他选项。')

H(doc, '2.1  为什么是 Node.js + Express？', 2)
FEAT(doc, '全栈语言统一',
     '前端使用 ArkTS（TypeScript 超集），后端使用 JavaScript。两者语法高度相似（ES6+ 的 async/await、箭头函数、解构赋值等），一个开发者可以无缝切换前后端，极大降低了全栈开发的心智负担。如果后端选 Python/Java，就需要在两种语言的思维模式间频繁切换。')
FEAT(doc, 'Express 的成熟度',
     'Express 是 Node.js 生态中 GitHub Stars 最多（60,000+）、npm 周下载量最大（3,000 万+）的 Web 框架。它的中间件模型（Middleware Pattern）是业界标准——app.use() 注册中间件，请求依次穿过每个中间件。这个模型直观、可测试、易扩展。')
FEAT(doc, '适合 RESTful API',
     'Express 的路由定义（app.get / app.post / app.put）天然映射 RESTful 语义。相比 Koa 的洋葱模型或 Nest.js 的装饰器模式，Express 的路由风格对于课设答辩来说"一眼就能看懂"，代码即文档。')
FEAT(doc, '零配置启动',
     '不需要像 Spring Boot 那样配置 Tomcat，不需要像 Django 那样运行 migrate。node server.js 一条命令即可启动，适合课设的快速演示。')

H(doc, '2.2  为什么是 JSON 文件数据库？', 2)
P(doc, '这是答辩时大概率会被问到的问题。以下是我们做出这个决策的完整逻辑：')
FEAT(doc, '极简部署',
     '不需要安装 MySQL、配置连接池、设置用户名密码。clone 代码 → npm install → node server.js，三秒启动。评审老师不需要在你的电脑上装数据库。')
FEAT(doc, '数据可视化',
     'db.json 可以用任何文本编辑器打开。答辩时你可以直接打开这个文件展示数据结构，比用 SQL 命令行 SELECT * FROM 直观得多。JSON 的层级结构一目了然。')
FEAT(doc, 'Git 友好',
     'db.json 是纯文本文件，天然支持 Git diff。每次修改数据后 git diff 可以看到具体哪些字段变化了。这在开发阶段非常有用。')
FEAT(doc, '代码即数据库',
     'readDB() 就是 JSON.parse(fs.readFileSync(...))，writeDB() 就是 JSON.stringify + fs.writeFileSync。数据库操作就是原生 JavaScript 对象操作——没有 ORM、没有 SQL 注入风险、没有连接池配置。')
FEAT(doc, '迁移成本低',
     '如果未来需要迁移到 MySQL/MongoDB，只需要改写 readDB() 和 writeDB() 两个函数。所有 API 处理函数使用的都是标准 JavaScript 对象，不感知底层存储。这是典型的"依赖倒置"设计。')

P(doc, '当然，我们清楚地认识到 JSON 文件数据库的局限性——不支持并发写入、没有事务回滚、大数据量下性能下降。但这些都是"工程权衡"而非"技术缺陷"。在课设阶段，数据量 < 1000 条、单用户操作、无需并发——这些条件完美匹配 JSON 文件方案。')

H(doc, '2.3  为什么只有两个依赖？', 2)
P(doc, 'package.json 中 dependencies 只有 express 和 cors。这是刻意为之的设计决策：')
B(doc, '身份认证：使用 Node.js 内置 crypto 模块（PBKDF2 + randomBytes），不用 jsonwebtoken/passport 等第三方库')
B(doc, '数据存储：使用 Node.js 内置 fs 模块（文件读写），不用 mongoose/sequelize 等 ORM')
B(doc, 'ID 生成：自研 genId() 函数（时间戳+随机数），不用 uuid 库')
B(doc, '安全收益：零第三方认证库依赖 = 零供应链攻击风险（还记得 left-pad 事件吗？）')
B(doc, '教学价值：从底层理解每一项技术，而不是"npm install 一个黑盒"')

PAGE(doc)

# ====================================================================
# 三、服务器架构逐层拆解
# ====================================================================
H(doc, '三、服务器架构逐层拆解', 1)

P(doc, '这是答辩的核心章节——你需要能够从宏观到微观完整描述整个后端的架构。')

H(doc, '3.1  架构总览（五层模型）', 2)
P(doc, '我们的后端采用了一个清晰的五层架构（从上到下）：')

P(doc, '【第 1 层】中间件层（Middleware Layer）', bold=True)
CODE(doc, 'app.use(cors())            // CORS 跨域处理')
CODE(doc, 'app.use(express.json())    // JSON 请求体解析')
CODE(doc, 'app.use(authMiddleware)    // Token 认证（自研）')
P(doc, '这三个中间件组成了请求的"安检通道"。每个请求在到达业务处理函数之前，必须依次通过这三道门。关键是：authMiddleware 是"软认证"——它只尝试解析 Token，解析不到不会拒绝请求，而是让 req.userId 保持 undefined，由下游业务函数自行判断。')

P(doc, '【第 2 层】路由层（Route Layer）', bold=True)
P(doc, 'Express 根据 HTTP 方法（GET/POST/PUT）和 URL 路径（/api/posts/:id）将请求分发到对应的处理函数。路由定义遵循 RESTful 规范——URL 使用资源名词（posts/users/comments），HTTP 方法表达操作语义（GET=读、POST=创建、PUT=更新）。')

P(doc, '【第 3 层】业务逻辑层（Business Logic Layer）', bold=True)
P(doc, '这是代码量最大的层。每个路由处理函数包含：参数校验 → 读取数据 → 业务判断 → 修改数据 → 写回文件 → 返回响应。以一个完整的"星尘打赏"操作为例，这一层执行了：余额校验 → 扣款 → 转账 → 等级升级判定 → 打赏记录 → 通知创建 → 数据持久化，共 7 个步骤。')

P(doc, '【第 4 层】数据访问层（Data Access Layer）', bold=True)
CODE(doc, 'function readDB() {')
CODE(doc, "  const raw = fs.readFileSync(DB_PATH, 'utf-8')")
CODE(doc, '  return JSON.parse(raw)')
CODE(doc, '}')
CODE(doc, 'function writeDB(data) {')
CODE(doc, "  fs.writeFileSync(DB_PATH, JSON.stringify(data, null, 2), 'utf-8')")
CODE(doc, '}')
P(doc, '只有这两个函数。它们封装了对 db.json 的全部 I/O 操作。如果未来换成 MySQL，只需要改这两个函数——上面三层的代码完全不用变。')

P(doc, '【第 5 层】基础设施层（Infrastructure Layer）', bold=True)
CODE(doc, 'genId(prefix)       // 唯一 ID 生成：前缀_时间戳36进制_随机4位')
CODE(doc, 'hashPassword(pw,salt) // PBKDF2-SHA512 密码哈希')
CODE(doc, 'makeSalt()          // crypto.randomBytes(16) 生成随机盐')
CODE(doc, 'generateToken()     // Token 生成：lm_时间戳_随机16位')
CODE(doc, 'setToken(userId)    // Token 存储到内存 Map')
CODE(doc, 'calculateLevel(totalReceived) // 等级计算')
CODE(doc, 'createNotification(...)       // 通知创建（所有通知的统一入口）')

H(doc, '3.2  请求处理完整时序（以"打赏帖子"为例）', 2)
P(doc, '下面追踪一个完整的 POST /api/posts/:id/reward 请求：')
CODE(doc, '① 客户端发送：POST /api/posts/p_001/reward')
CODE(doc, '   Body: { fromUserId: "u_001", amount: 10, message: "好文！" }')
CODE(doc, '   Header: Authorization: Bearer lm_xxx_yyy')
CODE(doc, '')
CODE(doc, '② CORS 中间件：放行（设置 Access-Control-Allow-Origin）')
CODE(doc, '③ JSON 中间件：req.body = { fromUserId, amount, message }')
CODE(doc, '④ 认证中间件：从 Header 提取 Token → tokenStore 查找 → req.userId = "u_001"')
CODE(doc, '⑤ 路由匹配：app.post("/api/posts/:id/reward", handler)')
CODE(doc, '⑥ 业务处理：')
CODE(doc, '   - readDB() 读取整个 db.json')
CODE(doc, '   - db.posts.find(p => p.id === "p_001") → 找到帖子')
CODE(doc, '   - db.users.find(u => u.id === "u_001") → 找到打赏者')
CODE(doc, '   - 检查 fromUser.starDustBalance >= 10 → 通过')
CODE(doc, '   - fromUser.starDustBalance -= 10  // 扣款')
CODE(doc, '   - toUser.totalStarDustReceived += 10  // 入账')
CODE(doc, '   - const lv = calculateLevel(toUser.totalStarDustReceived)')
CODE(doc, '   - toUser.level = lv.level; toUser.levelTitle = lv.levelTitle')
CODE(doc, '   - post.starDustCount += 10')
CODE(doc, '   - db.starDustRecords.unshift({记录对象})')
CODE(doc, '   - createNotification(db, "reward", ...)  // 通知作者')
CODE(doc, '   - writeDB(db)  // 持久化')
CODE(doc, '⑦ 返回：{ code: 0, data: 打赏记录, message: "打赏成功" }')

PAGE(doc)

# ====================================================================
# 四、数据模型完整设计
# ====================================================================
H(doc, '四、数据模型完整设计', 1)

P(doc, 'Lumina 的 JSON 数据库包含 8 个数据集合。以下逐一详解每个集合的设计意图、关键字段、和与业务逻辑的关联。')

# --- users ---
H(doc, '4.1  users（用户集合）—— 最核心的实体', 2)
P(doc, '用户是整个系统的中心节点。我们在此集合中内嵌了星尘余额、等级、社交计数等字段，避免高频查询时的"多表关联"。')
CODE(doc, '{')
CODE(doc, '  id: "u_001",                    // genId("u") 生成')
CODE(doc, '  studentId: "2021001",            // 学号，注册/登录凭证')
CODE(doc, '  nickname: "星尘旅人",              // 昵称，所有展示场景使用')
CODE(doc, '  avatar: "🐱",                    // 头像（emoji 或 URL）')
CODE(doc, '  bio: "在Lumina，每一个灵魂...",    // 个人简介')
CODE(doc, '  campus: "星光大学城",              // 校区')
CODE(doc, '  major: "计算机科学与技术",          // 专业')
CODE(doc, '  grade: "大三",                    // 年级')
CODE(doc, '  passwordHash: "c01b1841546...",  // PBKDF2-SHA512 64字节哈希')
CODE(doc, '  passwordSalt: "09a142c13cf3...", // 16字节随机盐')
CODE(doc, '  starDustBalance: 252,            // 当前星尘余额（可流通）')
CODE(doc, '  totalStarDustReceived: 1680,     // 累计收到星尘（用于等级判定）')
CODE(doc, '  postCount: 24,                   // 发帖总数（冗余计数）')
CODE(doc, '  followingCount: 56,              // 关注数（冗余计数）')
CODE(doc, '  followerCount: 129,              // 粉丝数（冗余计数）')
CODE(doc, '  level: 3,                        // 等级 1-6')
CODE(doc, '  levelTitle: "星辰旅者",            // 等级称号')
CODE(doc, '  joinDate: 1749600000000          // 注册时间戳')
CODE(doc, '}')
P(doc, '设计要点：postCount/followingCount/followerCount 是冗余字段。它们本可以通过 COUNT 关联表得到，但我们选择在用户表中冗余存储，因为"用户卡片"是最高频的展示场景。每次展示用户卡片都去 COUNT 所有帖子/关注/粉丝在数据量大时是性能灾难。"空间换时间"是后端优化的经典策略。')

# --- posts ---
H(doc, '4.2  posts（帖子集合）', 2)
CODE(doc, '{')
CODE(doc, '  id: "p_001",                     // genId("p") 生成')
CODE(doc, '  title: "期末复习资料大合集！",       // 标题')
CODE(doc, '  content: "整理了一份超全的...",     // 正文（支持长文本）')
CODE(doc, '  images: [],                      // 图片 URL 数组')
CODE(doc, '  authorId: "u_003",               // 作者 ID（外键）')
CODE(doc, '  authorName: "代码诗人",            // 作者名（冗余，避免 JOIN）')
CODE(doc, '  authorAvatar: "",                // 作者头像（冗余）')
CODE(doc, '  tag: "学习",                      // 分类标签（10种可选）')
CODE(doc, '  createTime: 1749744000000,       // 创建时间戳')
CODE(doc, '  likeCount: 329,                  // 点赞数')
CODE(doc, '  commentCount: 59,                // 评论数')
CODE(doc, '  starDustCount: 52,               // 累计收到星尘数')
CODE(doc, '  viewCount: 2912,                 // 浏览量')
CODE(doc, '  isTop: true,                     // 是否置顶')
CODE(doc, '  isHot: true,                     // 是否热门')
CODE(doc, '  likedBy: ["u_002","u_004",...]   // 点赞用户 ID 数组')
CODE(doc, '}')
P(doc, '设计要点：likedBy 数组替代了独立的"点赞表"。点赞时 push userId，取消点赞时 splice 移除。O(n) 的 includes 查找在课设数据量下完全可接受。authorName/authorAvatar 的冗余存储避免了每次渲染 PostCard 都去查用户表——这是高频场景下的性能优化。')

# --- comments ---
H(doc, '4.3  comments（评论集合）', 2)
CODE(doc, '{')
CODE(doc, '  id: "c_001",            // genId("c")')
CODE(doc, '  postId: "p_001",        // 所属帖子 ID（外键）')
CODE(doc, '  content: "太及时了！",   // 评论内容')
CODE(doc, '  authorId: "u_002",      // 评论者 ID')
CODE(doc, '  authorName: "图书馆小仙女"')
CODE(doc, '  createTime: 1749745000000')
CODE(doc, '  likeCount: 23')
CODE(doc, '  likedBy: ["u_001"]')
CODE(doc, '}')

# --- starDustRecords ---
H(doc, '4.4  starDustRecords（星尘流水记录）', 2)
P(doc, '这是星尘经济系统的"账本"。每一笔星尘流转——无论是签到获取、打赏支出、还是被打赏收入——都记录在此。用户可以在"我的星尘"页面查看完整的收支明细。')
CODE(doc, '{')
CODE(doc, '  id: "s_001"')
CODE(doc, '  fromUserId: "u_002",  fromUserName: "图书馆小仙女"')
CODE(doc, '  toUserId: "u_003",    toUserName: "代码诗人"')
CODE(doc, '  postId: "p_001",      postTitle: "期末复习资料大合集"')
CODE(doc, '  amount: 10,           message: "太有帮助了！"')
CODE(doc, '  createTime: 1749748000000')
CODE(doc, '}')

# --- notifications ---
H(doc, '4.5  notifications（通知集合）', 2)
P(doc, '通知系统支持 4 种类型：like（点赞通知）、comment（评论通知）、reward（打赏通知）、official（官方消息）。')
CODE(doc, '{')
CODE(doc, '  id: "n_001"')
CODE(doc, '  type: "like",          // like | comment | reward | official')
CODE(doc, '  fromUserId: "u_002",   fromUserName: "图书馆小仙女"')
CODE(doc, '  toUserId: "u_003",     // 接收者')
CODE(doc, '  postId: "p_001",       postTitle: "期末复习资料大合集"')
CODE(doc, '  content: "",           // official 类型有额外正文')
CODE(doc, '  isRead: false,         // 已读标记')
CODE(doc, '  createTime: 1749748000000')
CODE(doc, '}')

# --- follows ---
H(doc, '4.6  follows（关注关系）', 2)
CODE(doc, '{ fromUserId: "u_001", toUserId: "u_003", createTime: 1749600000000 }')
P(doc, '极简设计——只存三个字段。查询粉丝：follows.filter(f => f.toUserId === id)。查询关注：follows.filter(f => f.fromUserId === id)。')

# --- conversations & messages ---
H(doc, '4.7  conversations + messages（私信系统）', 2)
P(doc, '会话表存储对话元信息（参与者+最后消息），消息表存储具体消息。两个表通过 conversationId 关联。')
CODE(doc, '// conversations:')
CODE(doc, '{ id: "conv_001", user1: "u_001", user2: "u_003",')
CODE(doc, '  lastMessage: "Hello", lastTime: 1782323102046 }')
CODE(doc, '')
CODE(doc, '// messages:')
CODE(doc, '{ id: "msg_001", conversationId: "conv_001",')
CODE(doc, '  fromUserId: "u_001", toUserId: "u_003",')
CODE(doc, '  content: "Hello", createTime: 1782323102046, isRead: false }')

PAGE(doc)

# ====================================================================
# 五、认证与安全体系
# ====================================================================
H(doc, '五、认证与安全体系（答辩重点）', 1)

P(doc, '安全是后端开发的底线。Lumina 从零实现了完整的认证安全体系，未使用任何第三方认证库。')

H(doc, '5.1  密码存储：PBKDF2-SHA512 加盐哈希（完整流程）', 2)
P(doc, '这是答辩时最能体现技术深度的部分。以下是完整的密码处理流程：')

P(doc, '注册时：', bold=True)
B(doc, '① 前端发送：POST /api/auth/register { studentId, password, nickname }')
B(doc, '② 后端校验：密码长度 >= 6，学号未被注册')
B(doc, '③ 生成随机盐：const salt = crypto.randomBytes(16).toString("hex")')
B(doc, '   crypto.randomBytes 是 Node.js 内置的密码学安全随机数生成器')
B(doc, '   16 字节 = 128 位，生成空间为 2^128，碰撞概率接近零')
B(doc, '④ 密码加盐哈希：const hash = crypto.pbkdf2Sync(password, salt, 1000, 64, "sha512").toString("hex")')
B(doc, '   参数含义：password=原始密码, salt=随机盐, 1000=迭代次数, 64=输出64字节, sha512=哈希算法')
B(doc, '⑤ 存储：passwordHash = hash, passwordSalt = salt')
B(doc, '⑥ 原始密码被丢弃，服务器内存中不留痕迹')

P(doc, '登录时：', bold=True)
B(doc, '① 前端发送：POST /api/auth/login { studentId, password }')
B(doc, '② 后端根据 studentId 查找用户，取出该用户的 salt')
B(doc, '③ 对输入的密码执行相同哈希：hashPassword(inputPassword, user.passwordSalt)')
B(doc, '④ 比对：hash === user.passwordHash')
B(doc, '⑤ 匹配成功 → 生成 Token → 返回脱敏用户信息')
B(doc, '⑥ 匹配失败 → 返回"密码错误"')

P(doc, '安全核心：即使两个用户设置相同密码，由于盐值不同，哈希值完全不同。攻击者无法通过彩虹表反推。', bold=True, color=(0, 102, 255))

H(doc, '5.2  Token 认证机制（自研，不依赖 JWT）', 2)
P(doc, '我们选择自研 Token 方案而非 JWT，原因如下：')

FEAT(doc, '不受 JWT 无状态限制',
     'JWT 一旦签发就无法撤销（除非引入黑名单）。我们的 Token 方案中，服务端 tokenStore Map 存储 token→userId 映射，重启服务即可全局登出所有用户，管理员也可以随时删除特定 token 强制用户下线。')
FEAT(doc, 'Token 不含用户信息',
     'JWT 的 payload 是 Base64 编码（非加密！），任何人都可以解码看到用户 ID。我们的 Token 格式为 lm_{时间戳}_{随机字符串}，完全随机，不含任何用户信息。')
FEAT(doc, '单点登录',
     'setToken() 函数在生成新 Token 前会清除该用户的所有旧 Token：for (const [k, v] of tokenStore) { if (v === userId) tokenStore.delete(k); }。保证一个用户只有一个有效 Token。')

CODE(doc, '// Token 生成（server.js 第 43-45 行）')
CODE(doc, "function generateToken() {")
CODE(doc, "  return 'lm_' + Date.now().toString(36) + '_' + Math.random().toString(36).substr(2, 16)")
CODE(doc, "}")

H(doc, '5.3  数据脱敏', 2)
P(doc, '所有返回用户信息的 API 端点，都使用 ES6 解构赋值排除敏感字段：')
CODE(doc, 'const { passwordHash, passwordSalt, ...safeUser } = user')
CODE(doc, 'res.json({ code: 0, data: safeUser })')
P(doc, '这确保了 passwordHash 和 passwordSalt 绝对不会出现在任何 HTTP 响应中。即使前端开发者想获取这两个字段，也是不可能的。')

H(doc, '5.4  业务层面安全校验（6 道防线）', 2)
B(doc, '防线 1 - 密码长度：password.length < 6 → 400 "密码至少6位"')
B(doc, '防线 2 - 学号唯一性：注册时检查是否已存在 → "该学号已注册"')
B(doc, '防线 3 - 自关注防护：userId === targetId → "不能关注自己"')
B(doc, '防线 4 - 重复关注防护：follows 中已存在 → "已关注"')
B(doc, '防线 5 - 余额校验：starDustBalance < amount → "星尘余额不足"')
B(doc, '防线 6 - 自通知屏蔽：fromUserId === toUserId → 不创建通知')

PAGE(doc)

# ====================================================================
# 六、API 端点逐一详解
# ====================================================================
H(doc, '六、API 端点逐一详解（25 个接口）', 1)
P(doc, '以下是所有 25 个 API 端点的完整说明。每个端点标注了：HTTP 方法、URL 路径、请求参数、业务逻辑流程、返回值格式。你需要能口头讲清楚其中 8-10 个核心端点。')

# 认证
H(doc, '6.1  认证模块（3 个）', 2)

FEAT(doc, 'POST /api/auth/register',
     '入参 { studentId, password, nickname }。注册流程：校验学号唯一 → 校验密码>=6位 → makeSalt() → hashPassword() → 创建用户对象(初始100星尘, Lv.1) → db.users.push → writeDB → setToken → 返回 { token, user }。')
FEAT(doc, 'POST /api/auth/login',
     '入参 { studentId, password }。登录流程：学号查用户 → 取出 salt → hashPassword(输入密码, salt) → 比对 hash → setToken → 解构排除敏感字段 → 返回 { token, user }。')
FEAT(doc, 'GET /api/auth/me',
     '无需入参（从 authMiddleware 获取 req.userId）。直接查用户 → 脱敏返回。用于 App 启动时恢复登录状态。')

# 用户
H(doc, '6.2  用户模块（10 个）', 2)

FEAT(doc, 'GET /api/users/:id',
     '获取用户完整信息。URL 参数 :id 为用户 ID。返回包含等级、星尘余额、社交统计的完整用户对象。')
FEAT(doc, 'PUT /api/users/:id',
     '更新用户资料。入参 { nickname, bio, avatar, campus, major, grade }。支持部分更新——只传要改的字段，未传字段保持不变。')
FEAT(doc, 'POST /api/users/:id/checkin',
     '每日签到。无请求体。直接 user.starDustBalance += 5，返回最新余额。注意：当前版本未做"一天只能签到一次"的校验，这是后续迭代点。')
FEAT(doc, 'POST /api/users/:id/follow',
     '关注用户。入参 { userId }（关注者 ID）。流程：校验不关注自己 → 校验未重复关注 → db.follows.push → 双方计数+1 → writeDB。')
FEAT(doc, 'POST /api/users/:id/unfollow',
     '取消关注。入参 { userId }。从 follows 数组中 splice 移除 → 双方计数-1（Math.max(0, ...) 防止负数）。')
FEAT(doc, 'GET /api/users/:id/follow-status',
     '查询关注状态。入参 ?userId=当前用户。返回 { isFollowing: true/false }。用于进入他人主页时展示"已关注"/"关注"按钮。')
FEAT(doc, 'GET /api/users/:id/followers', '粉丝列表。返回 [{ id, nickname, avatar, bio, studentId }]。')
FEAT(doc, 'GET /api/users/:id/following', '关注列表。返回格式同上。')
FEAT(doc, 'GET /api/users/:id/posts', '用户帖子列表。按 createTime 降序排列。')
FEAT(doc, 'GET /api/users/:id/rewards', '星尘记录。双向匹配：fromUserId 或 toUserId 等于该用户 ID 的记录都返回。')

# 帖子
H(doc, '6.3  帖子模块（8 个）', 2)

FEAT(doc, 'GET /api/posts',
     '帖子列表。查询参数 ?tag=分类（可选）。排序规则：置顶优先（isTop 为 true 的排最前），然后按 createTime 降序。不传 tag 或 tag=全部时返回全站帖子。')
FEAT(doc, 'POST /api/posts',
     '发布帖子。入参 { title, content, authorId, tag, images }。流程：查作者 → 创建 post 对象 → db.posts.unshift(插到最前) → author.postCount+1 → author.starDustBalance+=2（发帖奖励）→ writeDB。')
FEAT(doc, 'GET /api/posts/hot',
     '热榜。核心算法：对全站帖子按热度分数降序排列。热度 = likeCount×2 + commentCount×3 + starDustCount×5 + viewCount×0.1。实时计算，无缓存。')
FEAT(doc, 'GET /api/posts/:id',
     '帖子详情。浏览量自动+1（post.viewCount++）。这是一个典型的"读操作附带副作用"设计。')
FEAT(doc, 'POST /api/posts/:id/like',
     '点赞/取消点赞 Toggle。入参 { userId }。检查 likedBy 是否包含 userId → 包含则移除(取消点赞)，不包含则添加(点赞) → createNotification → writeDB。返回 { isLiked, likeCount }。')
FEAT(doc, 'GET /api/posts/:id/comments',
     '评论列表。按 createTime 降序（最新的在前）。')
FEAT(doc, 'POST /api/posts/:id/comments',
     '发表评论。入参 { content, authorId }。流程：校验帖子存在 → 创建 comment → post.commentCount++ → createNotification → writeDB。')
FEAT(doc, 'POST /api/posts/:id/reward',
     '星尘打赏。入参 { fromUserId, amount, message }。这是最复杂的端点之一，包含 7 步操作：查帖子 → 查打赏者 → 余额校验 → 扣款 → 对方入账 → 等级升级判定 → 创建打赏记录 → createNotification → writeDB。')

# 通知
H(doc, '6.4  通知模块（3 个）', 2)
FEAT(doc, 'GET /api/notifications', '通知列表。查询参数 ?userId=。按 createTime 降序。前端按 type 字段分 Tab 展示。')
FEAT(doc, 'GET /api/notifications/unread-count', '未读计数。只返回一个整数。设计为独立端点是为了让前端可以高频轮询（只传输几个字节）。')
FEAT(doc, 'POST /api/notifications/read', '标记已读。入参 { userId, notificationId? }。传 notificationId → 标记单条已读。只传 userId → 全部已读。')

# 消息
H(doc, '6.5  消息模块（3 个）', 2)
FEAT(doc, 'POST /api/messages/send', '发送私信。自动查找或创建会话 → 创建消息 → 更新 lastMessage/lastTime → writeDB。')
FEAT(doc, 'GET /api/messages/conversations', '会话列表。返回每个会话的对方用户信息 + 最后消息 + 未读计数。')
FEAT(doc, 'GET /api/messages/:convId', '会话消息列表。按 createTime 正序（旧→新）。自动将 toUserId === userId 的消息标记为已读。')

# 搜索 + 管理
H(doc, '6.6  搜索 + 管理模块（4 个）', 2)
FEAT(doc, 'GET /api/search?q=关键字', '全字段搜索。将帖子标题、正文、标签 toLowerCase() 后进行 includes 匹配。简单的模糊搜索，课设数据量下完全够用。')
FEAT(doc, 'POST /api/admin/bootstrap', '一键初始化演示数据。包含：创建6个粉丝 → 全部关注目标用户 → 对帖子点赞评论 → 3条官方消息 → 种子帖子。')
FEAT(doc, 'POST /api/admin/broadcast', '全站广播。入参 { title, content }。遍历所有用户，为每人创建一条 official 类型通知。')
FEAT(doc, 'POST /api/admin/seed', '种子数据。生成 8 篇知识型帖子（《人类简史》读后感、量子计算、康德哲学、算法竞赛等），标题和正文都是真实内容。')

PAGE(doc)

# ====================================================================
# 七、核心业务逻辑剖析
# ====================================================================
H(doc, '七、核心业务逻辑剖析', 1)

H(doc, '7.1  打赏系统的完整执行路径', 2)
P(doc, '打赏是 Lumina 最复杂的业务操作。以下是 POST /api/posts/:id/reward 的完整伪代码：')
CODE(doc, '1.  const db = readDB()')
CODE(doc, '2.  const post = db.posts.find(p => p.id === req.params.id)')
CODE(doc, '3.  if (!post) → 404 "帖子不存在"')
CODE(doc, '4.  const fromUser = db.users.find(u => u.id === fromUserId)')
CODE(doc, '5.  if (!fromUser) → 404 "用户不存在"')
CODE(doc, '6.  if (fromUser.starDustBalance < amount) → 400 "余额不足"')
CODE(doc, '7.  fromUser.starDustBalance -= amount          // ① 扣款')
CODE(doc, '8.  const toUser = db.users.find(u => u.id === post.authorId)')
CODE(doc, '9.  toUser.totalStarDustReceived += amount       // ② 入账')
CODE(doc, '10. const lv = calculateLevel(toUser.totalStarDustReceived)')
CODE(doc, '11. toUser.level = lv.level; toUser.levelTitle = lv.levelTitle  // ③ 升级')
CODE(doc, '12. post.starDustCount += amount                  // ④ 帖子计数')
CODE(doc, '13. db.starDustRecords.unshift(record)             // ⑤ 流水记录')
CODE(doc, '14. createNotification(db, "reward", ...)          // ⑥ 通知')
CODE(doc, '15. writeDB(db)                                   // ⑦ 持久化')
CODE(doc, '16. res.json({ code: 0, data: record, message: "打赏成功" })')

H(doc, '7.2  六等级体系', 2)
P(doc, '用户等级由累计收到的星尘总数决定，在 calculateLevel() 函数中判定：')
CODE(doc, 'function calculateLevel(totalReceived) {')
CODE(doc, '  if (totalReceived >= 10000) return { level: 6, levelTitle: "星海领主" }')
CODE(doc, '  if (totalReceived >= 5000)  return { level: 5, levelTitle: "星尘大师" }')
CODE(doc, '  if (totalReceived >= 2000)  return { level: 4, levelTitle: "星光使者" }')
CODE(doc, '  if (totalReceived >= 500)   return { level: 3, levelTitle: "星辰旅者" }')
CODE(doc, '  if (totalReceived >= 100)   return { level: 2, levelTitle: "星尘学徒" }')
CODE(doc, '  return { level: 1, levelTitle: "星尘新芽" }')
CODE(doc, '}')
P(doc, '等级升级发生在两个场景：① 被打赏时（reward 端点中）② 发帖时（post 端点中，收到打赏后）。升级是即时判定的——每次 starDustReceived 增加后立即调用 calculateLevel。')

H(doc, '7.3  热榜算法公式', 2)
P(doc, '热度 = 点赞×2 + 评论×3 + 星尘×5 + 浏览×0.1', bold=True, size=12, color=(0, 102, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
P(doc, '')
P(doc, '每项权重的设计依据：星尘(5x) > 评论(3x) > 点赞(2x) > 浏览(0.1x)。星尘权重最高是因为打赏者付出了真实虚拟资产，代表了最强的认可信号。浏览权重极低是为了防止"标题党"——即使标题吸引大量点击，如果内容不好（没人点赞/评论/打赏），也不会出现在热榜高位。')
P(doc, '这个算法在每次 GET /api/posts/hot 时实时计算。代码实现：')
CODE(doc, 'const hotA = a.likeCount*2 + a.commentCount*3 + a.starDustCount*5 + a.viewCount*0.1')
CODE(doc, 'const hotB = b.likeCount*2 + b.commentCount*3 + b.starDustCount*5 + b.viewCount*0.1')
CODE(doc, 'return hotB - hotA  // Array.sort 降序')

H(doc, '7.4  通知生成机制', 2)
P(doc, 'createNotification() 是所有通知的唯一入口函数。它被以下 4 个场景调用：')
B(doc, '点赞时 → createNotification(db, "like", likerId, likerName, post.authorId, post.id, post.title)')
B(doc, '评论时 → createNotification(db, "comment", commenterId, commenterName, post.authorId, post.id, post.title)')
B(doc, '打赏时 → createNotification(db, "reward", fromUserId, fromUserName, post.authorId, post.id, post.title)')
B(doc, '广播时 → createNotification(db, "official", "admin", "Lumina官方", eachUser.id, "", title)')

P(doc, '核心设计：函数第一行就检查 fromUserId === toUserId，如果给自己点赞/评论/打赏，直接 return（不产生通知）。这个细节必须在答辩时提到——体现了对用户体验的细致思考。')

H(doc, '7.5  私信系统的会话管理', 2)
P(doc, '发送消息时，系统自动处理会话的创建或更新：')
B(doc, '① 查找是否已存在这两个用户之间的会话（双向匹配：user1+user2 或 user2+user1）')
B(doc, '② 如果不存在 → 创建新会话（genId("conv")）')
B(doc, '③ 创建消息对象，关联到会话 ID')
B(doc, '④ 更新会话的 lastMessage（截取前 50 字符）和 lastTime')

PAGE(doc)

# ====================================================================
# 八、管理后台与演示数据系统
# ====================================================================
H(doc, '八、管理后台与演示数据系统', 1)

P(doc, '这是 Lumina 课设的一大亮点——专门为答辩演示设计了管理 API。')

H(doc, '8.1  /api/admin/seed —— 知识型种子帖子', 2)
P(doc, '这个端点生成 8 篇高质量种子帖子，涵盖读书分享、科技前沿、人文社科、课程交流、职业发展、学术讨论等分类。每篇帖子有完整的标题和正文——内容摘录自真实学术讨论话题，如尤瓦尔·赫拉利的《人类简史》、Google IBM 量子计算进展、康德与黑格尔比较、MIT 6.824 课程笔记等。')
P(doc, '这些帖子的存在，使得演示时 App 看起来像一个"已经运营了一段时间的知识社区"，而不是一个空白的 Demo。')

H(doc, '8.2  /api/admin/bootstrap —— 一键生成演示数据', 2)
P(doc, '这个端点可以在 1 次 API 调用中完成以下全部操作：')
B(doc, '创建 6 个粉丝账号（算法爱好者 / 哲学系小张 / 量子萌新 / 读书人小王 / 职场观察者 / 思辨者），每个有独立的昵称、简介、人格设定')
B(doc, '所有粉丝自动关注目标用户（学号 123456）')
B(doc, '粉丝对用户的每篇帖子执行点赞 + 评论（6 种评论话术轮换使用）')
B(doc, '生成 36+ 条通知（点赞通知 + 评论通知）')
B(doc, '广播 3 条官方消息给全体用户')
B(doc, '一次性展示"消息爆炸"效果')

P(doc, '答辩演示时建议的操作顺序：先在 App 中注册账号(学号123456) → 调用 bootstrap → 刷新 App → 展示通知列表中的 36 条未读 → 展示粉丝数从 0 变成 6 → 展示帖子收到的点赞和评论。', bold=True, color=(0, 102, 255))

H(doc, '8.3  /api/admin/broadcast —— 全站广播', 2)
P(doc, '管理员通过此端点向所有注册用户推送官方消息。技术实现是遍历 db.users，为每个用户创建一条 official 类型通知。入参 { title, content }，返回 { sentCount }。')

PAGE(doc)

# ====================================================================
# 九、部署与运维
# ====================================================================
H(doc, '九、部署与运维指南', 1)

H(doc, '9.1  环境要求', 2)
B(doc, 'Node.js 16.0 及以上版本')
B(doc, 'npm（随 Node.js 一起安装）')
B(doc, '任何操作系统（Windows / macOS / Linux）')

H(doc, '9.2  启动步骤', 2)
CODE(doc, 'cd C:/Lumina/server')
CODE(doc, 'npm install       # 首次运行需要，之后不需要')
CODE(doc, 'node server.js    # 启动服务')
P(doc, '启动成功后会看到：')
CODE(doc, '✨ Lumina API Server running at http://localhost:3000')
CODE(doc, '   POST   /api/auth/login       - 登录')
CODE(doc, '   POST   /api/auth/register    - 注册')
CODE(doc, '   ... (全部 25 个端点列表)')

H(doc, '9.3  模拟器联调', 2)
P(doc, 'HarmonyOS 模拟器中，10.0.2.2 是 Android 模拟器的特殊 IP，映射到宿主机的 localhost。前端的 ApiClient.ets 中 BASE_URL 默认配置为此地址。')

H(doc, '9.4  真机联调', 2)
B(doc, '确保手机和电脑在同一 WiFi 网络')
B(doc, '查看电脑的局域网 IP：ipconfig（Windows）→ 找到 IPv4 地址，如 192.168.1.100')
B(doc, '修改 ApiClient.ets 中的 BASE_URL 为 http://192.168.1.100:3000')
B(doc, '关闭 Windows 防火墙对 3000 端口的拦截（或添加入站规则）')

H(doc, '9.5  测试 API 的方法', 2)
P(doc, '推荐使用以下工具测试后端 API：')
B(doc, 'Postman（GUI 工具，最直观，适合演示）')
B(doc, 'curl 命令行（如 curl http://localhost:3000/api/posts）')
B(doc, '浏览器直接访问 GET 端点（如 http://localhost:3000/api/posts）')

PAGE(doc)

# ====================================================================
# 十、答辩 QA 预判
# ====================================================================
H(doc, '十、答辩 QA 预判', 1)

P(doc, '以下整理了答辩时评审老师最可能问的问题，以及参考答案。', bold=True)

FEAT(doc, 'Q1：为什么用 JSON 文件而不用 MySQL？',
     '答：这是经过深思熟虑的工程决策。课设阶段数据量小（<1000条）、单用户操作、无并发需求——这些条件完美匹配 JSON 文件方案。同时，我们将数据访问封装在 readDB/writeDB 两个函数中，未来只需修改这两个函数即可迁移到任何数据库。这是"依赖倒置原则"的实践。')
FEAT(doc, 'Q2：后端安全性如何保证？',
     '答：我们实现了三层安全防护。① 密码安全：PBKDF2-SHA512 + 随机盐 + 1000次迭代哈希，业界标准级别。② 传输安全：Bearer Token 认证机制，Token 不含用户信息。③ 数据安全：所有返回用户信息的 API 使用解构赋值排除密码字段。此外还有 6 道业务级别的安全校验。')
FEAT(doc, 'Q3：如果用户量增长到 10 万，系统哪里会先出问题？',
     '答：三个瓶颈。① JSON 文件读写——每次请求都全量读取整个 db.json，数据量大会导致 IO 延迟。解决方案：迁移到 SQLite/MySQL。② 热榜实时计算——全站帖子排序，帖子多了 O(n log n) 耗时增加。解决方案：Redis 缓存 + 定时任务预计算。③ likedBy 数组的 includes 查找——O(n) 复杂度。解决方案：改用 Set 数据结构或独立点赞表。')
FEAT(doc, 'Q4：Token 方案和 JWT 相比有什么优劣？',
     '答：我们没有使用 JWT 而是自研了 Token。优势：Token 可随时撤销（服务端 Map 中删除即可）；Token 不含用户信息，即使泄露也不会暴露用户 ID；不依赖第三方库。劣势：需要服务端存储状态，不支持水平扩展（但课设场景无此需求）。')
FEAT(doc, 'Q5：热榜算法的权重为什么这样设计？',
     '答：星尘(5x) > 评论(3x) > 点赞(2x) > 浏览(0.1x)。这个权重反映了不同互动行为的"价值层级"——打赏需要付出真实虚拟资产，是最强的认可信号；评论需要文字输入，参与度高于点赞；浏览权重极低是为了防止标题党。这个设计体现了"用经济学手段引导用户行为"的产品思维。')
FEAT(doc, 'Q6：通知系统为什么用轮询而不用 WebSocket？',
     '答：当前实现是基于 HTTP 轮询的。我们为未读计数设计了独立的轻量级端点（只返回一个整数），将轮询开销降到最低。课设阶段轮询完全够用。如果未来需要实时推送，HarmonyOS 端有 @ohos.net.webSocket API 可以直接对接，后端增加 socket.io 即可。')
FEAT(doc, 'Q7：你在这个项目中最自豪的技术设计是什么？',
     '答：两个。一是星尘经济系统的闭环设计——从注册送星尘、签到获取、打赏消费、到等级晋升，形成了一个完整的激励正反馈循环。二是一键演示数据系统——bootstrap 端点可以在 1 次 API 调用中生成全部演示数据，这体现了"为演示而设计"的工程思维，在课设答辩中非常实用。')

P(doc, '')
P(doc, '')
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('— 后端技术深度解读报告 · 完 —')
rr.font.size = Pt(13); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.save('Lumina_后端技术深度解读.docx')
print('Done: 后端技术深度解读.docx')


# ============================================================
#           报告二：小组分工报告
# ============================================================
doc2 = setup_doc()

cover(doc2, 'Lumina 项目', '小组分工与贡献说明',
      ['—— 知识型学术社区 · 产品设计与技术实现',
       '课程：XXXX 课程设计',
       '指导老师：XXX'])

H(doc2, '一、项目概述', 1)
P(doc2, 'Lumina 是一款面向高校学生的知识型学术社区 App，基于 HarmonyOS 5.0（ArkTS）开发前端，Node.js + Express 构建后端 RESTful API 服务。项目涵盖产品设计、UI/UX 设计、前端开发、后端开发、测试联调等完整技术链路，由 6 名组员协作完成。')
P(doc2, '项目整体产出：前端 20 个页面（约 4,500 行 ArkTS 代码）、后端 25 个 API 端点（约 860 行 Node.js 代码）、数据库 8 个数据集合（JSON 文件存储）、自研双主题颜色系统（约 150 行）、星尘虚拟货币激励体系（完整经济闭环）。以下为各成员的具体分工与贡献。')

H(doc2, '二、小组成员与分工明细', 1)

# ---- 分工表格（6人） ----
table = doc2.add_table(rows=7, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['姓名', '负责模块', '具体工作内容', '技术难度']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for pp in cell.paragraphs:
        for run in pp.runs:
            run.bold = True; run.font.size = Pt(10)

# 权重递减：葛沛林 > 潘家乐 > 陆喆 > 程雅婷 > 史浩然 > 任一诺
divisions = [
    ['葛沛林\n（队长）',
     '后端系统架构\n（核心引擎）',
     '① 整体后端技术选型与五层架构设计（中间件→路由→业务→数据→基础设施）\n'
     '② 全部 25 个 RESTful API 端点的编码实现（认证/用户/帖子/评论/星尘/通知/关注/私信/搜索/管理共 10 大模块）\n'
     '③ 8 个数据集合的建模与 JSON 数据库设计（users/posts/comments/starDustRecords/notifications/follows/conversations/messages）\n'
     '④ PBKDF2-SHA512 密码加盐哈希 + Bearer Token 认证安全体系（自研，零第三方认证库依赖）\n'
     '⑤ 星尘经济系统后端核心逻辑（打赏转账 / 余额校验 / 等级自动升级判定 / 流水记录）\n'
     '⑥ 热榜加权算法设计与实现（热度 = 点赞×2 + 评论×3 + 星尘×5 + 浏览×0.1）\n'
     '⑦ 通知推送引擎（createNotification 统一入口，支持 like/comment/reward/official 四种类型）\n'
     '⑧ 私信系统后端（会话自动创建/消息收发/未读标记）\n'
     '⑨ 管理后台 API（一键演示数据 bootstrap / 种子帖子 seed / 全站广播 broadcast / 批量粉丝生成）\n'
     '⑩ 前后端接口规范制定（统一 { code, data, message } 响应格式 + Token 注入约定）',
     '★★★★★'],

    ['潘家乐',
     '前端核心页面\n+ 社交消息系统',
     '① 首页信息流（HomePage）：帖子列表渲染 + 10 种标签横向滚动筛选栏 + 下拉刷新\n'
     '② 热榜页面（HotPage）：时间维度切换（今日/本周/本月/全部）+ 金/银/铜排名徽章\n'
     '③ 帖子详情页（PostDetailPage）：正文富文本展示 + 点赞（红心动画）+ 评论列表\n'
     '④ 消息通知页面（MessagePage）：5 类通知 Tab（全部/赞/评论/官方/私信）+ 全部已读\n'
     '⑤ 私信聊天系统（ChatListPage + ChatPage）：会话列表 + 一对一聊天 + 未读角标\n'
     '⑥ PostCard 可复用帖子卡片组件开发（被首页/热榜/搜索结果/用户帖子等多处复用）\n'
     '⑦ 底部导航栏 Index 页面架构（5 个 Tab：首页/热榜/[发布]/消息/我的）+ 中间突出发布按钮',
     '★★★★'],

    ['陆喆',
     '星尘激励系统\n+ 个人中心',
     '① 星尘经济系统的产品机制设计（获取途径：签到/发帖/被打赏；消费：打赏他人；等级晋升规则）\n'
     '② 星尘打赏面板（StarDustRewardPage）：6 档金额选择（1/5/10/20/50/100）+ 自定义留言 + 动画效果\n'
     '③ 星尘规则页面（StarDustRulesPage）：完整展示获取途径 + 等级体系说明\n'
     '④ 个人中心页面（ProfilePage）：用户信息卡片 + 数据统计（帖子/粉丝/关注）+ 功能菜单\n'
     '⑤ 星尘明细底部弹出面板（余额展示 + 收支流水列表 + 等级进度）\n'
     '⑥ 首页星尘签到横幅组件（签到+5 星尘 + 签到后按钮变灰防重复点击）\n'
     '⑦ 六等级徽章 UI 组件（星尘新芽→星尘学徒→星辰旅者→星光使者→星尘大师→星海领主）\n'
     '⑧ 编辑资料页面（EditProfilePage）：头像选择 + 信息修改 + 保存刷新',
     '★★★★'],

    ['程雅婷',
     'UI/UX 设计\n+ 账号认证模块',
     '① App 整体视觉风格设计（知乎风格参考：简洁、克制、信息优先）\n'
     '② 深色/浅色双主题色彩体系设计（浅色 7 色 + 深色 7 色，共 16 个语义化颜色属性）\n'
     '③ Constants.ets 主题管理器实现（ThemeManager + C 颜色代理类，AppStorage 驱动全局切换）\n'
     '④ 字体/间距/圆角设计规范制定（统一 Spacing / BorderRadius / FontSize 常量）\n'
     '⑤ Logo 与品牌视觉元素设计 + 启动闪屏（SplashPage）渐入动画\n'
     '⑥ 登录页面（LoginPage）+ 注册页面（RegisterPage）前端实现\n'
     '⑦ 设置页面（SettingsPage）：主题切换 + 退出登录 + 关于入口\n'
     '⑧ 用户资料展示页（UserProfilePage）：查看他人主页 + 关注按钮',
     '★★★'],

    ['史浩然',
     '内容发布与搜索\n+ 联调测试',
     '① 帖子发布页面（PublishPage）：标题输入 + 正文编辑 + 10 种知识标签选择 + 最多 9 张图片选择\n'
     '② 搜索页面（SearchPage）：搜索栏 + 全字段模糊搜索结果列表 + 热门搜索词快捷填充\n'
     '③ ApiClient 前端网络层封装（GET/POST/PUT 方法 + 自动 Token 注入 + 超时配置 + 错误处理）\n'
     '④ PostApi / UserApi 业务 API 封装（帖子 CRUD + 用户操作 + 关注/取关）\n'
     '⑤ Model 层 JSON → 类实例安全映射实现（PostModel / UserModel / CommentModel / StarDustModel）\n'
     '⑥ 全功能联调测试（前后端 25 个 API 逐一验证 + 20 个页面功能走查）\n'
     '⑦ 模拟器与真机兼容性验证（Phone + Tablet 双形态适配）',
     '★★★'],

    ['任一诺',
     '辅助页面与组件\n+ 关注社交模块',
     '① 关注/粉丝列表页面（FollowListPage）：列表渲染 + 关注状态切换 + 跳转用户主页\n'
     '② 关于页面（AboutPage）：App 信息展示 + 版本号 + 技术栈说明\n'
     '③ 路由配置与页面导航逻辑（20 个页面的 router.push 跳转 + 参数传递）\n'
     '④ 全局状态管理实现（AppStorage 的 loginVersion / feedNeedsRefresh + @Watch 回调机制）\n'
     '⑤ CurrentUser 全局单例设计（登录后用户信息跨页面共享，避免重复请求）\n'
     '⑥ 空数据状态友好提示组件 + 未登录 Toast 拦截逻辑\n'
     '⑦ 协助潘家乐完成首页信息流 + 热榜页面的部分 UI 开发工作',
     '★★'],
]

for r_idx, row_data in enumerate(divisions):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(8.5)

doc2.add_paragraph()
P(doc2, '注：葛沛林同学作为项目队长，独立承担了后端全栈开发这一技术难度最高、工作量最大的核心环节。后端 860 行代码从架构设计到全部 25 个 API 编码实现均为独立完成，并主导了前后端接口规范制定与安全体系建设。', bold=True, size=10, color=(0, 102, 255))

H(doc2, '三、工作量占比估算', 1)
P(doc2, '以下按各成员实际承担的技术难度与代码贡献量估算工作量占比（权重递减：葛沛林 > 潘家乐 > 陆喆 > 程雅婷 > 史浩然 > 任一诺）：')

table3 = doc2.add_table(rows=7, cols=4, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['排名', '姓名', '主要贡献领域', '工作量占比']
for i, h_text in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h_text
    for pp in cell.paragraphs:
        for run in pp.runs:
            run.bold = True; run.font.size = Pt(10)

work_data = [
    ['1', '葛沛林（队长）', '后端全栈（25 API + 安全体系 + 数据库 + 星尘引擎 + 管理后台）', '30%'],
    ['2', '潘家乐', '前端核心页面（首页/热榜/详情/消息/私信/Index 导航）+ PostCard 组件', '22%'],
    ['3', '陆喆', '星尘激励系统（产品设计 + 打赏面板 + 个人中心 + 等级体系 + 签到）', '18%'],
    ['4', '程雅婷', 'UI/UX 设计 + 双主题系统 + 账号认证模块（登录/注册/设置/用户主页）', '14%'],
    ['5', '史浩然', '发布页 + 搜索页 + ApiClient 网络层 + Model 映射 + 联调测试', '10%'],
    ['6', '任一诺', '关注/粉丝/关于页面 + 路由导航 + 状态管理 + 协助首页开发', '6%'],
]
for r_idx, row_data in enumerate(work_data):
    for c_idx, cell_text in enumerate(row_data):
        cell = table3.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(10)

doc2.add_paragraph()

H(doc2, '四、协作方式', 1)
B(doc2, '版本控制：使用 Git 进行代码版本管理，葛沛林（队长）负责代码合并、冲突解决与最终 Review')
B(doc2, '前后端接口约定：葛沛林定义全部 25 个 API 接口规范（URL 路径、请求参数、响应格式），前端同学按接口文档并行开发')
B(doc2, '进度同步：每周进行一次进度对齐会议，使用微信群进行日常沟通与问题协调')
B(doc2, '代码审查：后端全部代码 + 前端 ApiClient 网络层由葛沛林 Review 把关；前端页面代码组员间交叉 Review')
B(doc2, '测试联调：史浩然牵头编写测试用例，全体成员在各自模块完成后参与端到端联调')

doc2.add_paragraph()
doc2.add_paragraph()
pp = doc2.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('— 小组分工报告 · 完 —')
rr.font.size = Pt(13); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc2.save('Lumina_小组分工报告.docx')
print('Done: 小组分工报告.docx')
