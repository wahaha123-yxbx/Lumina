# -*- coding: utf-8 -*-
"""
生成两份文档：
  1. Lumina_小组分工报告.docx  —— 修正版（去掉权重%和难度★）
  2. Lumina_课设答辩汇报书.docx —— 课设风格，用于汇报展示
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os

os.chdir('C:/Lumina')

# ============================================================
#                       工具函数
# ============================================================
def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'; style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)
    return doc

def H(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def P(doc, text, bold=False, color=None, size=11):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor(*color)
    pp.paragraph_format.space_after = Pt(4)
    pp.paragraph_format.line_spacing = 1.5

def B(doc, text):
    pp = doc.add_paragraph(text, style='List Bullet')
    pp.paragraph_format.space_after = Pt(2); pp.paragraph_format.line_spacing = 1.4
    for run in pp.runs:
        run.font.size = Pt(10.5); run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def CODE(doc, text):
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(1); pp.paragraph_format.line_spacing = 1.2
    r = pp.add_run('  ' + text)
    r.font.name = 'Consolas'; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(60, 60, 60)

def KEY(doc, title, content):
    pp = doc.add_paragraph()
    pp.paragraph_format.space_after = Pt(4); pp.paragraph_format.line_spacing = 1.5
    r1 = pp.add_run('📍 ' + title + '：')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGBColor(0, 102, 255)
    r1.font.name = '微软雅黑'; r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r2 = pp.add_run(content)
    r2.font.size = Pt(10.5); r2.font.name = '微软雅黑'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def PAGE(doc):
    doc.add_page_break()

def cover(doc, title, subtitle, desc_lines):
    for _ in range(4): doc.add_paragraph()
    t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run(title)
    r.font.size = Pt(44); r.bold = True; r.font.color.rgb = RGBColor(0, 102, 255)
    r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    t1b = doc.add_paragraph(); t1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1b = t1b.add_run(subtitle)
    r1b.font.size = Pt(15); r1b.font.color.rgb = RGBColor(100, 100, 100)
    r1b.font.name = '微软雅黑'; r1b.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph(); doc.add_paragraph()
    today = datetime.date.today().strftime('%Y 年 %m 月 %d 日')
    for ml in desc_lines + [today]:
        tt = doc.add_paragraph(); tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tt.add_run(ml).font.size = Pt(10.5)
    PAGE(doc)


# ================================================================
#        报告一：小组分工报告（修正版——无%无★）
# ================================================================
doc1 = setup_doc()
cover(doc1, 'Lumina 项目', '小组分工与贡献说明',
      ['知识型学术社区 · 产品设计与技术实现',
       '课程：XXXX 课程设计    |    指导老师：XXX'])

H(doc1, '一、项目概述', 1)
P(doc1, 'Lumina 是一款面向高校学生的知识型学术社区 App，基于 HarmonyOS 5.0（ArkTS）开发前端，Node.js + Express 构建后端 RESTful API 服务。项目涵盖产品设计、UI/UX 设计、前端开发、后端开发、测试联调等完整技术链路，由 6 名组员协作完成。')
P(doc1, '项目整体产出：前端 20 个页面（约 4,500 行 ArkTS 代码）、后端 25 个 API 端点（约 860 行 Node.js 代码）、数据库 8 个数据集合（JSON 文件存储）、自研双主题颜色系统、星尘虚拟货币激励体系。')

H(doc1, '二、小组成员与分工明细', 1)

table = doc1.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['姓名', '负责模块', '具体工作内容']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(10)

divisions = [
    ['葛沛林\n（队长）',
     '后端系统架构',
     '① 整体后端技术选型与五层架构设计\n'
     '② 全部 25 个 RESTful API 端点的编码实现（认证/用户/帖子/评论/星尘/通知/关注/私信/搜索/管理共 10 大模块）\n'
     '③ 8 个数据集合的建模与 JSON 数据库设计\n'
     '④ PBKDF2-SHA512 密码加盐哈希 + Bearer Token 认证安全体系\n'
     '⑤ 星尘经济系统后端核心逻辑（打赏/转账/等级升级判定/流水记录）\n'
     '⑥ 热榜加权算法设计与实现\n'
     '⑦ 通知推送引擎（统一入口，4 种通知类型）\n'
     '⑧ 私信系统后端（会话管理/消息收发/未读标记）\n'
     '⑨ 管理后台 API（一键演示数据/种子帖子/全站广播）\n'
     '⑩ 前后端接口规范制定（统一响应格式 + Token 注入约定）'],

    ['潘家乐',
     '前端核心页面\n+ 社交消息系统',
     '① 首页信息流（HomePage）+ 10 种标签横向滚动筛选栏\n'
     '② 热榜页面（HotPage）+ 时间维度切换 + 金/银/铜排名徽章\n'
     '③ 帖子详情页（PostDetailPage）+ 点赞动画 + 评论列表展示\n'
     '④ 消息通知页面（MessagePage）+ 5 类通知 Tab + 全部已读\n'
     '⑤ 私信聊天系统（ChatListPage + ChatPage）+ 未读角标\n'
     '⑥ PostCard 可复用帖子卡片组件（多页面复用）\n'
     '⑦ 底部导航栏 Index 页面架构（5 个 Tab + 中间突出发布按钮）'],

    ['陆喆',
     '星尘激励系统\n+ 个人中心',
     '① 星尘经济系统的产品机制设计（获取途径/消费方式/等级体系）\n'
     '② 星尘打赏面板 + 6 档金额选择 + 自定义留言 + 动画效果\n'
     '③ 星尘规则说明页面\n'
     '④ 个人中心页面（ProfilePage）+ 数据统计 + 功能菜单\n'
     '⑤ 星尘明细底部弹出面板（余额/收支流水/等级进度）\n'
     '⑥ 首页星尘签到横幅组件\n'
     '⑦ 六等级徽章 UI 组件\n'
     '⑧ 编辑资料页面（EditProfilePage）'],

    ['程雅婷',
     'UI/UX 设计\n+ 账号认证模块',
     '① App 整体视觉风格设计（知乎风格：简洁/克制/信息优先）\n'
     '② 深色/浅色双主题色彩体系 + Constants.ets 主题管理器实现\n'
     '③ 字体/间距/圆角设计规范制定\n'
     '④ Logo 与品牌视觉元素 + 启动闪屏渐入动画\n'
     '⑤ 登录页面 + 注册页面前端实现\n'
     '⑥ 设置页面（主题切换/退出登录）\n'
     '⑦ 用户资料展示页（UserProfilePage）+ 关注按钮'],

    ['史浩然',
     '内容发布与搜索\n+ 网络层与联调',
     '① 帖子发布页面（PublishPage）+ 知识标签选择 + 图片选择\n'
     '② 搜索页面 + 全字段模糊搜索 + 热门搜索词\n'
     '③ ApiClient 前端网络层封装（GET/POST/PUT + Token 注入 + 超时配置）\n'
     '④ PostApi / UserApi 业务 API 封装\n'
     '⑤ Model 层 JSON → 类实例安全映射实现\n'
     '⑥ 全功能联调测试（前后端 25 个 API 逐一验证）\n'
     '⑦ 模拟器 + 真机兼容性验证'],

    ['任一诺',
     '辅助页面与组件\n+ 关注社交模块',
     '① 关注/粉丝列表页面 + 关注状态切换\n'
     '② 关于页面\n'
     '③ 路由配置与页面导航逻辑（20 个页面路由跳转 + 参数传递）\n'
     '④ 全局状态管理实现（loginVersion/feedNeedsRefresh + @Watch）\n'
     '⑤ CurrentUser 全局单例设计\n'
     '⑥ 空数据友好提示组件 + 未登录 Toast 拦截\n'
     '⑦ 协助潘家乐完成首页 + 热榜页面部分 UI 开发'],
]

for r_idx, row_data in enumerate(divisions):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(8.5)

doc1.add_paragraph()

H(doc1, '三、协作方式', 1)
B(doc1, '版本控制：使用 Git 进行代码版本管理，葛沛林负责代码合并、冲突解决与最终 Review')
B(doc1, '前后端接口约定：葛沛林定义全部 25 个 API 接口规范（URL 路径、请求参数、响应格式），前端同学按接口文档并行开发')
B(doc1, '进度同步：每周进行进度对齐，使用微信群日常沟通')
B(doc1, '代码审查：后端全部代码 + 前端 ApiClient 网络层由葛沛林 Review 把关；前端页面代码组员间交叉 Review')
B(doc1, '测试联调：史浩然牵头编写测试用例，全体成员参与端到端联调')

doc1.add_paragraph()
pp = doc1.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('— 小组分工报告 · 完 —')
rr.font.size = Pt(13); rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc1.save('Lumina_Team_Division.docx')
print('Done: Team Division')


# ================================================================
#        报告二：课设答辩汇报书
# ================================================================
doc2 = setup_doc()
cover(doc2, 'Lumina', '知识型学术社区 · 课设答辩汇报',
      ['基于 HarmonyOS + Node.js 的全栈 App 开发',
       '汇报人：葛沛林    |    小组：程雅婷 / 任一诺 / 陆喆 / 潘家乐 / 史浩然'])

# ===== 目录 =====
H(doc2, '目  录', 1)
toc = ['一、项目背景与定位', '二、产品功能演示路线', '三、技术架构总览',
       '四、后端系统设计（重点）', '五、前端架构设计', '六、创新特色功能',
       '七、项目成果统计', '八、总结与展望']
for t in toc: P(doc2, t)
PAGE(doc2)

# ===== 一、项目背景 =====
H(doc2, '一、项目背景与定位', 1)

H(doc2, '1.1  出发点', 2)
P(doc2, '高校学生面临一个普遍矛盾：渴望获取有深度的知识内容，但主流社交平台的内容越来越碎片化和娱乐化。有价值的学术讨论和深度见解被淹没在短视频和信息流中。')
P(doc2, '我们的答案：Lumina —— 一款专为高校学生打造的知识型学术社区。不是问答工具，不是娱乐平台，而是一个以知识分享为纽带的学术社交空间。')

H(doc2, '1.2  产品定位', 2)
B(doc2, 'Slogan：「以知识为星尘，照亮彼此的视野」')
B(doc2, '目标用户：在校本科生及研究生（18-26 岁）')
B(doc2, '核心差异：自研「星尘」虚拟货币体系驱动内容质量正循环，区别于传统点赞驱动的社交平台')
B(doc2, '运行平台：华为 HarmonyOS 5.0 (API 12+)，兼容 Phone + Tablet')

H(doc2, '1.3  与竞品的差异', 2)
table = doc2.add_table(rows=4, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['', '传统社交平台', '传统问答平台', 'Lumina']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)
for r, row in enumerate([
    ['内容调性', '娱乐化、碎片化', '工具化、冷冰冰', '知识型、有温度的学术社区'],
    ['激励方式', '点赞数/粉丝数', '积分/徽章', '星尘虚拟货币 + 六等级声望体系'],
    ['用户关系', '单向关注', '匿名问答', '双向关注 + 知识打赏经济闭环'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

PAGE(doc2)

# ===== 二、产品功能演示路线 =====
H(doc2, '二、产品功能演示路线', 1)
P(doc2, '以下是答辩演示的推荐流程，约 5-8 分钟可完整展示全部核心功能：')

B(doc2, '① 启动 App → 品牌闪屏 → 首页信息流（展示标签筛选 + 帖子列表）')
B(doc2, '② 深色/浅色主题一键切换（展示自研双主题系统）')
B(doc2, '③ 热榜页面（展示加权热度排序 + 金/银/铜排名徽章）')
B(doc2, '④ 帖子详情 → 点赞（红心动画）→ 星尘打赏（弹出面板选择金额）→ 评论')
B(doc2, '⑤ 发布新帖子（标题/正文/标签/图片选择）')
B(doc2, '⑥ 搜索功能（输入关键词 → 全字段模糊搜索）')
B(doc2, '⑦ 消息通知（展示分类 Tab + 一键全部已读）')
B(doc2, '⑧ 私信聊天（会话列表 → 一对一对话）')
B(doc2, '⑨ 个人中心（用户数据 + 星尘明细 + 等级徽章 + 关注/粉丝）')
B(doc2, '⑩ 调用 /api/admin/bootstrap 一键生成演示数据 → 展示通知爆炸效果')

PAGE(doc2)

# ===== 三、技术架构总览 =====
H(doc2, '三、技术架构总览', 1)

P(doc2, 'Lumina 采用前后端分离架构，前端为 HarmonyOS 原生 App，后端为 Node.js RESTful API 服务。')

H(doc2, '3.1  技术栈一览', 2)
table = doc2.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['', '前端', '后端']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)
for r, row in enumerate([
    ['开发语言', 'ArkTS（TypeScript 超集）', 'JavaScript (Node.js)'],
    ['框架/平台', 'ArkUI 声明式 UI + HarmonyOS SDK', 'Express 4.18'],
    ['关键依赖', 'HarmonyOS 5.0 API 12', '仅 express + cors（2 个依赖）'],
    ['数据存储', 'AppStorage（状态管理）', 'JSON 文件数据库（8 个集合）'],
    ['构建工具', 'Hvigor 5.0 + DevEco Studio 6.0', 'npm + node 直接运行'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

doc2.add_paragraph()

H(doc2, '3.2  前后端通信方式', 2)
B(doc2, '协议：HTTP JSON，统一响应格式 { code: 0, data: ..., message: ... }')
B(doc2, '认证：Bearer Token（自研方案，非 JWT），Header 携带 Authorization: Bearer <token>')
B(doc2, '开发环境：模拟器通过 10.0.2.2:3000 访问宿主机后端；真机通过局域网 IP 访问')

PAGE(doc2)

# ===== 四、后端系统设计（重点） =====
H(doc2, '四、后端系统设计（重点）', 1)

H(doc2, '4.1  技术选型理由', 2)
KEY(doc2, 'Node.js + Express',
    '前端 ArkTS 是 TypeScript 超集，后端用 JavaScript——全栈语言统一。Express 是 Node.js 生态最成熟的 Web 框架，中间件模型直观，路由定义清晰。')
KEY(doc2, 'JSON 文件数据库',
    '课设数据量 < 1000 条，JSON 完全满足。优势：零安装配置、Git 版本控制友好、文本编辑器可直接查看。数据访问封装在 readDB/writeDB 两个函数中，未来迁移数据库只需改这两个函数。')
KEY(doc2, '仅 2 个依赖（express + cors）',
    '密码哈希用 Node.js 内置 crypto（PBKDF2），Token 自研不用 JWT，ID 自研不用 uuid。零第三方认证库 = 零供应链风险，同时展示从底层理解安全机制的工程能力。')

H(doc2, '4.2  五层架构设计', 2)
B(doc2, '第 1 层 — 中间件链：CORS 跨域 → JSON 解析 → authMiddleware 认证（软认证：解析不到 Token 不拒绝请求）')
B(doc2, '第 2 层 — 路由分发：25 个端点按 RESTful 语义分发（GET=读、POST=创建、PUT=更新）')
B(doc2, '第 3 层 — 业务逻辑：参数校验 → readDB → 业务判断 → 修改 → writeDB → 响应')
B(doc2, '第 4 层 — 数据访问：readDB() / writeDB() 两个函数封装全部文件 IO（依赖倒置原则）')
B(doc2, '第 5 层 — 基础设施：genId / hashPassword / generateToken / calculateLevel / createNotification')

H(doc2, '4.3  数据模型（8 个集合）', 2)
table = doc2.add_table(rows=9, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['集合', '核心说明']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)
for r, row in enumerate([
    ['users（用户）', '含密码哈希+盐值、星尘余额/累计、等级、社交计数（冗余字段，避免高频COUNT）'],
    ['posts（帖子）', '含作者信息冗余、likedBy[]替代独立点赞表、10种知识分类标签'],
    ['comments（评论）', '通过 postId 关联帖子，likedBy 同样用数组存储'],
    ['starDustRecords', '星尘流水账本，记录每笔流转的双方/金额/留言'],
    ['notifications', '4 种类型 like|comment|reward|official，isRead 标记已读/未读'],
    ['follows', '极简二元组 (fromUserId, toUserId)，双向查询粉丝/关注'],
    ['conversations', '会话元信息，lastMessage 冗余最后消息预览'],
    ['messages', '通过 conversationId 关联会话，按时间正序展示'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9)

H(doc2, '4.4  安全体系设计', 2)

KEY(doc2, '密码安全 —— PBKDF2-SHA512 加盐哈希',
    '注册时为每个用户生成 16 字节独立随机盐（crypto.randomBytes），使用 PBKDF2-SHA512 进行 1000 次迭代哈希。核心优势：即使两个用户设置相同密码，由于盐值不同，数据库中的哈希完全不一样。攻击者无法使用彩虹表。')
KEY(doc2, '认证安全 —— 自研 Token 方案',
    '不使用 JWT，自研 Token（格式 lm_xxx_xxx）。优势：Token 纯随机不含用户信息、服务端可随时撤销、单点登录（新 Token 自动清除旧 Token）。存储于内存 Map<token, userId>。')
KEY(doc2, '数据安全 —— 响应脱敏',
    '所有返回用户信息的接口使用解构赋值排除密码字段：const { passwordHash, passwordSalt, ...safeUser } = user。')
KEY(doc2, '业务安全 —— 6 道校验',
    '密码长度校验 | 学号唯一性 | 自关注防护 | 重复关注防护 | 打赏余额校验 | 自通知屏蔽')

H(doc2, '4.5  API 接口统计', 2)
table = doc2.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['模块', '端点数量与说明']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(9.5)
for r, row in enumerate([
    ['认证模块', '3 个：注册 (PBKDF2哈希) / 登录 (Token签发) / 获取当前用户'],
    ['用户模块', '10 个：信息CRUD / 签到 / 关注&取关 / 粉丝&关注列表 / 帖子&星尘记录'],
    ['帖子模块', '8 个：列表(标签筛选) / 发布(+2星尘) / 热榜(加权算法) / 详情(+浏览量) / 点赞Toggle / 评论CRUD / 打赏(7步事务)'],
    ['通知模块', '3 个：列表 / 未读计数(轻量轮询) / 标记已读'],
    ['消息模块', '3 个：发送 / 会话列表 / 会话消息(自动已读)'],
    ['搜索模块', '1 个：全字段模糊匹配(标题+正文+标签)'],
    ['管理模块', '4 个：种子数据 / 一键演示 / 全站广播 / 批量粉丝'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(9.5)

PAGE(doc2)

# ===== 五、前端架构设计 =====
H(doc2, '五、前端架构设计', 1)

H(doc2, '5.1  项目结构', 2)
B(doc2, '开发平台：HarmonyOS 5.0 (API 12) + ArkTS + ArkUI 声明式开发')
B(doc2, '页面规模：20 个 @Entry 页面 + 1 个 @Component 可复用组件（PostCard）')
B(doc2, '代码量：约 4,500 行 ArkTS 代码')

H(doc2, '5.2  六层架构', 2)
B(doc2, 'Entry 层：EntryAbility.ets — App 生命周期管理 + 窗口配置')
B(doc2, 'Page 层：20 个页面组件，@Entry 装饰，各自独立路由')
B(doc2, 'Component 层：PostCard 可复用卡片，被首页/热榜/搜索等多处引用')
B(doc2, 'Service 层：ApiClient（HTTP 封装 + Token 自动注入）+ PostApi + UserApi')
B(doc2, 'Model 层：JSON → 类实例安全映射（PostModel/UserModel/CommentModel/StarDustModel）')
B(doc2, 'Theme 层：ThemeManager + C 颜色代理类（约 150 行，深色/浅色双主题）')

H(doc2, '5.3  关键技术点', 2)
KEY(doc2, '状态管理',
    'loginVersion（登录版本号，@StorageLink + @Watch 实现跨组件同步刷新）+ feedNeedsRefresh（列表刷新标记）+ CurrentUser 全局单例')
KEY(doc2, '双主题系统',
    'C 类（颜色代理）：16 个 getter 属性，根据 ThemeManager.isDark() 返回对应色值。组件中写 C.bg / C.textPrimary，切换主题时自动重绘。')
KEY(doc2, 'JSON 安全映射',
    '后端返回纯 JSON（无类方法），前端每个 API 调用后执行显式 new Model() + 逐字段赋值，防止调用不存在的方法导致崩溃。')

PAGE(doc2)

# ===== 六、创新特色功能 =====
H(doc2, '六、创新特色功能', 1)

H(doc2, '6.1  星尘虚拟货币经济系统', 2)
P(doc2, '星尘是 Lumina 最核心的创新设计——一个闭环的虚拟货币激励体系：')
B(doc2, '获取途径：每日签到（+5）、发布帖子（+2）、被他人打赏')
B(doc2, '使用途径：打赏优质帖子（6 档金额：1/5/10/20/50/100，可附留言）')
B(doc2, '六等级体系：星尘新芽(Lv.1) → 星尘学徒(Lv.2, 累计100) → 星辰旅者(Lv.3, 500) → 星光使者(Lv.4, 2000) → 星尘大师(Lv.5, 5000) → 星海领主(Lv.6, 10000)')
B(doc2, '经济闭环：用户产出优质内容 → 被打赏获得星尘 → 用星尘打赏他人 → 激励更多优质内容')
P(doc2, '这套机制将内容质量的评价权从中心化"算法推荐"交还给去中心化的"社区投票"，同时解决了社区类产品的冷启动问题（新用户注册即送 100 星尘）。')

H(doc2, '6.2  热榜加权算法', 2)
P(doc2, '热度 = 点赞×2 + 评论×3 + 星尘×5 + 浏览×0.1', bold=True, size=13, color=(0, 102, 255))
P(doc2, '权重设计逻辑：星尘(5x) 代表最强认可（付出了真实资产）> 评论(3x) 代表深度参与（文字输入成本）> 点赞(2x) 代表轻度认可 > 浏览(0.1x) 防止标题党。每次请求热榜时对全站帖子实时计算。')

H(doc2, '6.3  一键演示数据系统', 2)
P(doc2, '/api/admin/bootstrap 端点可以在 1 次 API 调用中自动生成完整演示场景：')
B(doc2, '创建 6 个拟人化粉丝账号（算法爱好者 / 哲学系小张 / 量子萌新 / 读书人小王 / 职场观察者 / 思辨者）')
B(doc2, '所有粉丝自动关注目标用户 + 对帖子点赞评论（6 种评论话术轮换）')
B(doc2, '生成 36+ 条各类通知 + 3 条官方消息全站广播')
P(doc2, '这个功能让答辩演示时可以一键展示"消息爆炸"效果，无需手动逐条创建数据。')

PAGE(doc2)

# ===== 七、项目成果统计 =====
H(doc2, '七、项目成果统计', 1)

table = doc2.add_table(rows=9, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['指标', '数据']):
    cell = table.rows[0].cells[i]; cell.text = h
    for pp in cell.paragraphs:
        for run in pp.runs: run.bold = True; run.font.size = Pt(10)
for r, row in enumerate([
    ['前端页面数', '20 个独立页面 + 1 个可复用组件'],
    ['前端代码量', '约 4,500 行 ArkTS'],
    ['后端 API 端点数', '25 个 RESTful 端点，分 10 个业务模块'],
    ['后端代码量', '约 860 行 Node.js (server.js)'],
    ['数据集合数', '8 个（JSON 文件存储）'],
    ['核心依赖数', '仅 2 个（express + cors）'],
    ['支持设备类型', 'Phone + Tablet (HarmonyOS 5.0 / API 12+)'],
    ['安全机制', 'PBKDF2-SHA512 加盐哈希 + 自研 Token + 数据脱敏 + 6 道业务校验'],
]):
    for c, text in enumerate(row):
        cell = table.rows[r+1].cells[c]; cell.text = text
        for pp in cell.paragraphs:
            for run in pp.runs: run.font.size = Pt(10)

doc2.add_paragraph()

KEY(doc2, '技术栈全景',
    '后端：Node.js → Express → RESTful API → JSON DB → PBKDF2 密码哈希 → Bearer Token → 通知引擎')
KEY(doc2, '',
    '前端：HarmonyOS 5.0 → ArkTS → ArkUI 声明式 UI → AppStorage 状态管理 → 组件化架构 → 双主题系统')

PAGE(doc2)

# ===== 八、总结与展望 =====
H(doc2, '八、总结与展望', 1)

H(doc2, '8.1  项目总结', 2)
P(doc2, 'Lumina 项目从零到一完整构建了一款知识型社区 App 的全部前后端功能。我们实现了：')
B(doc2, '完整的前后端分离架构 —— 20 个前端页面 + 25 个后端 API')
B(doc2, '自研安全认证体系 —— 从密码哈希到 Token 管理全部自主实现，零第三方认证库依赖')
B(doc2, '创新的星尘经济系统 —— 完整的虚拟货币闭环 + 六等级声望体系')
B(doc2, '工程化的代码组织 —— 五层后端架构 + 六层前端架构，层次清晰')
B(doc2, '为演示而设计的工程思维 —— 一键生成演示数据，答辩展示效果好')

H(doc2, '8.2  个人技术收获（葛沛林）', 2)
B(doc2, '深入掌握了 Node.js + Express 后端开发全流程 —— 从架构设计到 API 编码到安全加固')
B(doc2, '实践了 PBKDF2 密码哈希、Token 认证、数据脱敏等后端安全技术')
B(doc2, '积累了 RESTful API 设计经验 —— 25 个端点的统一规范设计')
B(doc2, '理解了 HarmonyOS ArkTS + ArkUI 声明式开发范式')
B(doc2, '学会了从产品视角思考技术方案 —— 星尘经济系统和一键演示系统体现了"技术为产品服务"的思维')

H(doc2, '8.3  未来改进方向', 2)
B(doc2, '数据库升级：JSON 文件 → SQLite（支持事务和并发控制）')
B(doc2, '图片上传：当前图片为 URL 文本 → 升级为 multipart/form-data 真实上传')
B(doc2, '实时通知：HTTP 轮询 → WebSocket 推送')
B(doc2, '内容审核：接入 AI 文本审核，对帖子和评论自动合规检测')
B(doc2, '推荐算法：基于标签和用户行为的协同过滤推荐')
B(doc2, '鸿蒙能力：利用分布式数据管理实现同账号多设备同步')

doc2.add_paragraph()
doc2.add_paragraph()
pp = doc2.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pp.add_run('感谢各位老师！欢迎提问。')
rr.font.size = Pt(16); rr.bold = True; rr.font.color.rgb = RGBColor(0, 102, 255)
rr.font.name = '微软雅黑'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc2.save('Lumina_Presentation_Report.docx')
print('Done: Presentation Report')
